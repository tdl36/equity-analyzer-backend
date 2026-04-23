#!/usr/bin/env python3
"""Charlie Local Agent -- processes note generation jobs using local iCloud files.

Runs on the user's Mac, polls the Charlie backend for pending note jobs,
reads source files from iCloud, calls Claude API, generates full research
notes with charts, and saves outputs locally + uploads to Charlie.

Usage:
    python3 charlie_local_agent.py                 # Run agent daemon
    python3 charlie_local_agent.py --api-key sk-...# Run with explicit API key
    python3 charlie_local_agent.py --once           # Process one job and exit
    python3 charlie_local_agent.py --ticker MDT     # Process specific ticker (no backend job needed)
    python3 charlie_local_agent.py --debug          # Verbose logging
"""

import os
import sys
import json
import base64
import io
import re
import time
import shutil
import signal
import argparse
import logging
import threading
import uuid
import zipfile
from pathlib import Path
from datetime import datetime
from typing import Optional
from xml.etree import ElementTree

import requests
import anthropic

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

STOCKS_DIR = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/STOCKS"
CHARLIE_API = "https://equity-analyzer-backend.onrender.com"
POLL_INTERVAL = 5  # seconds
CONFIG_FILE = Path.home() / ".charlie_agent_config.json"
AGENT_SECRET_HEADER = "X-Agent-Secret"

# Telegram notifications
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "8487228930:AAHqadJIzFEVGq5TUESrKufD8-nHxcufj7Y")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID", "465648202")

log = logging.getLogger("charlie_agent")


def notify(message: str) -> None:
    """Send a Telegram notification to the user."""
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        return
    try:
        requests.post(
            f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage",
            json={"chat_id": TELEGRAM_CHAT_ID, "text": message, "parse_mode": "Markdown"},
            timeout=10,
        )
    except Exception:
        pass  # Don't let notification failures break the agent

# ---------------------------------------------------------------------------
# Sector map (mirrors backend PIPELINE_SECTOR_MAP / PIPELINE_TICKER_SECTOR)
# ---------------------------------------------------------------------------

PIPELINE_SECTOR_MAP = {
    "Pharma": ["MRK", "LLY", "JNJ", "PFE", "BMY", "ABBV"],
    "Biotech": ["GILD", "REGN", "VRTX"],
    "MedTech & Life Sci": ["MDT", "BDX", "BSX", "ABT", "TMO", "DHR", "DGX"],
    "Managed Care & Distro": ["UNH", "CI", "CVS", "COR", "CAH"],
    "Industrials": ["RTX", "GD", "CARR", "ETN", "DE", "PH", "DOV", "MMM", "RSG", "UNP", "NSC"],
    "REITs": ["VTR", "PLD", "AVB", "AMT"],
}

PIPELINE_TICKER_SECTOR: dict[str, str] = {}
for _sec, _tks in PIPELINE_SECTOR_MAP.items():
    for _tk in _tks:
        PIPELINE_TICKER_SECTOR[_tk] = _sec


def get_sector(ticker: str) -> str:
    return PIPELINE_TICKER_SECTOR.get(ticker, "Other")


# ---------------------------------------------------------------------------
# Playbook (matches backend RESEARCH_NOTE_PLAYBOOK exactly)
# ---------------------------------------------------------------------------

RESEARCH_NOTE_PLAYBOOK = """
## NOTE FORMAT — FOLLOW THIS STRUCTURE EXACTLY

### Header (lines 1-4):
Line 1: # Company Name (TICKER)
Line 2: Month Year | Version X.X
Line 3: (blank)
Line 4: Price: ~$XX | Mkt Cap: ~$XXB | FY26E EPS: ~$X.XX | P/E: ~XXx | FCF Yield: ~X.X%
Line 5: Conclusion: Own / Avoid / Revisit at $X | Next Catalyst: [specific event + date]

### Section structure (numbered, use ## N. Title format):

## 1. Executive Summary / Investment Thesis
- Start with a 1-paragraph business overview (what the company does, key context)
- Then **Bull case (3-4 bullets):** with bold headers per bullet
- Then **Bear case (3-4 bullets):** with bold headers per bullet
- Then a **Probability-Weighted Scenario Analysis** markdown table:
  | Scenario | FY27E EPS | Target Price | Implied P/E | Probability | Key Assumptions |

## 2. Business Overview & Segment Breakdown
- Segment breakdown as a markdown table: | Segment | Revenue | % of Total | Adj OP | Margin |
- Reference charts inline: ![TICKER Revenue Breakdown](TICKER_Revenue_Breakdown.png)
- Key sub-business callouts with bold headers

## 3. Key Revenue & Earnings Drivers
- FY bridge with specific numbers (volume/mix, price, cost, commodity, capital allocation)
- Growth algorithm if management provides one
- Use sub-headers for different driver categories

## 4. What the Street Is Debating
- 2-3 numbered debates as sub-headers: **Debate #1: [question]**
- For each: state both sides, then "My view:" or "I think:" with a clear stance
- Be specific with data points supporting each side

## 5. Financial Summary
- Multi-year markdown table: | Metric | FY2023A | FY2024A | FY2025A | FY2026E | FY2027E |
- Include: Revenue, YoY Growth, Adj OP, OP Margin, Adj EPS, EPS Growth, FCF, Net Debt/EBITDA, Capital Returned

## 6. Valuation & Catalysts
- Current trading multiples (P/E, EV/EBITDA)
- Peer context with specific comparisons
- "My take:" paragraph with clear assessment
- Catalyst Calendar as markdown table: | Date | Event | Significance |

## 7. Management & Capital Allocation
- CEO background and track record
- Key strategic decisions and execution assessment
- Capital allocation priorities (buybacks, dividends, M&A, deleveraging)

## 8. Risks
- 5-7 bullet points, each with bold header and specific quantification where possible

## 9. Bottom Line
- 2-3 paragraphs with clear actionable conclusion
- Specific price levels for action
- Conditions to watch for thesis change
- Relative positioning vs other names in coverage

### RECENCY & DATA PRIORITY:
- When source documents span multiple reporting periods (e.g., Q3 vs Q4 earnings), ALWAYS use the most recent quarter's data for financial metrics, guidance, estimates, and forward targets
- Explicitly state which quarter/period the data reflects (e.g., "as of Q4 2025 earnings")
- If prior analysis referenced older data (e.g., Q3 guidance), update ALL relevant metrics to the latest available (revenue, EPS, margins, guidance ranges)
- Flag material changes from prior period: "Updated from Q3: revenue guidance raised from $X to $Y"
- Do NOT carry forward stale estimates when newer data is available in the source documents
- For valuation multiples and price targets, use the most recent earnings/guidance as the basis

### STRICT RULES:
- NEVER reference specific analyst names, firms, or broker ratings in the main note
- Synthesize data points but attribute nothing to specific brokers
- Hard facts (reported financials, FDA approvals, deals, guidance) = state directly
- Do NOT use sellside sentiment, consensus ratings, or target ranges as reasons to buy/sell
- The investment thesis must stand on fundamentals alone
- NEVER attribute headline YoY EPS growth to a single narrative driver without decomposing the bridge
- Cross-check every narrative claim against data tables
- When FY EPS includes >$0.50/share in non-recurring items, flag explicitly
- Distinguish reported vs underlying/organic growth rates
- Nudge displayed numbers +/- $200-300M from model to avoid exact replication
- Use markdown tables for ALL tabular data (segments, financials, catalysts, scenarios)

### SECTOR-SPECIFIC ADDITIONS:
- Pharma/Biotech: Patent cliffs, pipeline table, LOE timeline, stacked revenue charts
- MedTech: Procedure volume trends, ASP dynamics, new product cycles
- Managed Care: Membership trends, MLR, PBM reform risk, star ratings
- Distribution: Drug pricing dynamics, biosimilar opportunity, generic deflation
- Industrials: Cycle positioning, book-to-bill, aftermarket mix, margin expansion
- REITs: Same-store NOI, occupancy, cap rates, lease spreads, FFO/AFFO

### TONE:
- Write as if Tony (the analyst) is authoring the note to his PM
- Confident, concise, first-person: "I think...", "My view is..."
- No hedging language like "it could potentially maybe..."
- Lead with conclusion, support with evidence
- Use precise numbers — don't round $4.69B to "about $5B"
- No emojis, no filler, no fluff
"""

# ---------------------------------------------------------------------------
# Configuration helpers
# ---------------------------------------------------------------------------


def load_config() -> dict:
    """Load persisted config from ~/.charlie_agent_config.json."""
    if CONFIG_FILE.exists():
        try:
            return json.loads(CONFIG_FILE.read_text())
        except Exception:
            pass
    return {}


def save_config(cfg: dict) -> None:
    """Persist config to ~/.charlie_agent_config.json."""
    CONFIG_FILE.write_text(json.dumps(cfg, indent=2))


def load_api_key_from_config() -> Optional[str]:
    return load_config().get("api_key")


def get_agent_secret() -> Optional[str]:
    """Return agent secret for authenticating with backend agent endpoints."""
    return os.environ.get("CHARLIE_AGENT_SECRET") or load_config().get("agent_secret")


# ---------------------------------------------------------------------------
# Backend communication
# ---------------------------------------------------------------------------


def _agent_headers() -> dict:
    secret = get_agent_secret()
    headers = {"Content-Type": "application/json"}
    if secret:
        headers[AGENT_SECRET_HEADER] = secret
    # API key auth for backend authentication
    api_key = os.environ.get("CHARLIE_API_KEY", "")
    if api_key:
        headers["Authorization"] = f"ApiKey {api_key}"
    return headers


def poll_for_jobs() -> list[dict]:
    """Poll backend for pending note generation jobs."""
    try:
        resp = requests.get(
            f"{CHARLIE_API}/api/pipeline/jobs",
            params={"status": "queued", "limit": 10},
            headers=_agent_headers(),
            timeout=15,
        )
        if resp.status_code == 200:
            jobs = resp.json().get("jobs", [])
            # Filter for jobs the local agent handles
            return [j for j in jobs if j.get("job_type") in ("note", "synthesis", "scan_catalysts")]
    except requests.exceptions.ConnectionError:
        log.warning("Backend unreachable -- will retry")
    except Exception as e:
        log.warning(f"Poll error: {e}")
    return []


def claim_job(job_id: str) -> bool:
    """Attempt to claim a job via the agent claim endpoint.

    Falls back to direct progress update if the agent endpoint does not exist.
    """
    try:
        resp = requests.post(
            f"{CHARLIE_API}/api/agent/claim",
            json={"jobId": job_id, "agentId": _agent_id()},
            headers=_agent_headers(),
            timeout=15,
        )
        if resp.status_code == 200:
            return True
        # Fallback: use the existing progress update mechanism
        log.debug(f"Claim endpoint returned {resp.status_code}, falling back to progress update")
    except Exception as e:
        log.debug(f"Claim endpoint unavailable ({e}), using fallback")

    # Fallback: update job directly
    return _update_job_progress_fallback(job_id, "running", "Claimed by local agent", 1)


def update_job_progress(
    job_id: str,
    status: str,
    step: str,
    progress: Optional[int],
    error: Optional[str] = None,
    result: Optional[dict] = None,
) -> None:
    """Update pipeline job progress on Charlie backend."""
    try:
        payload = {
                "jobId": job_id,
                "status": status,
                "currentStep": step,
                "progress": progress,
                "error": error,
        }
        if result:
            payload["result"] = result
        resp = requests.post(
            f"{CHARLIE_API}/api/agent/update-job",
            json=payload,
            headers=_agent_headers(),
            timeout=15,
        )
        if resp.status_code != 200:
            # Fallback
            _update_job_progress_fallback(job_id, status, step, progress, error)
    except Exception:
        _update_job_progress_fallback(job_id, status, step, progress, error)


def _update_job_progress_fallback(
    job_id: str,
    status: str,
    step: str,
    progress: Optional[int],
    error: Optional[str] = None,
) -> bool:
    """Fallback: POST to a generic update endpoint or accept that we can't update."""
    # If no agent endpoints exist yet, log locally and continue
    log.debug(f"Job {job_id[:12]}... => {status} / {step} / {progress}%")
    return True


def complete_job(job_id: str, result: dict) -> None:
    """Mark a job as complete and upload result summary."""
    try:
        requests.post(
            f"{CHARLIE_API}/api/agent/complete",
            json={"jobId": job_id, "result": result},
            headers=_agent_headers(),
            timeout=30,
        )
    except Exception:
        pass
    # Always update status as well
    update_job_progress(job_id, "complete", "Complete", 100)


def upload_note_to_charlie(
    ticker: str,
    note_md: str,
    sources_md: str,
    changelog_md: str,
    docx_path: Optional[Path],
    chart_paths: list[Path],
    version: str,
    metadata: dict,
) -> bool:
    """Upload the completed note to Charlie's backend."""
    docx_b64 = ""
    if docx_path and docx_path.exists():
        docx_b64 = base64.b64encode(docx_path.read_bytes()).decode("ascii")

    charts_payload: list[dict] = []
    for cp in chart_paths:
        if cp and cp.exists():
            chart_type = "revenue" if "Revenue" in cp.name else "profit"
            charts_payload.append(
                {
                    "type": chart_type,
                    "filename": cp.name,
                    "data": base64.b64encode(cp.read_bytes()).decode("ascii"),
                }
            )

    note_id = str(uuid.uuid4())
    payload = {
        "noteId": note_id,
        "ticker": ticker,
        "version": version,
        "noteMarkdown": note_md,
        "sourcesMarkdown": sources_md,
        "changelogMarkdown": changelog_md,
        "noteDocx": docx_b64,
        "charts": [{"type": c["type"], "filename": c["filename"]} for c in charts_payload],
        "chartsData": charts_payload,
        "metadata": metadata,
    }

    try:
        resp = requests.post(
            f"{CHARLIE_API}/api/agent/save-note",
            json=payload,
            headers=_agent_headers(),
            timeout=120,
        )
        if resp.status_code == 200:
            log.info(f"Note uploaded to Charlie: {ticker} v{version}")
            return True
        else:
            log.warning(f"Upload returned {resp.status_code}: {resp.text[:200]}")
            return False
    except Exception as e:
        log.error(f"Upload failed: {e}")
        return False


def _agent_id() -> str:
    """Stable agent ID for this machine."""
    import platform

    return f"local-{platform.node()}"


def process_pending_syncs() -> None:
    """Check for pending sync-to-local jobs and write edited notes back to iCloud."""
    try:
        resp = requests.get(
            f"{CHARLIE_API}/api/agent/sync-to-local",
            headers=_agent_headers(),
            timeout=15,
        )
        if resp.status_code != 200:
            return
        syncs = resp.json().get("syncs", [])
        for sync in syncs:
            ticker = sync.get("ticker", "")
            if not ticker:
                continue
            ticker_dir = STOCKS_DIR / ticker
            if not ticker_dir.exists():
                log.warning(f"  Sync: no folder for {ticker}, skipping")
                continue
            month = datetime.now().strftime("%b%Y")
            # Write markdown files
            note_md = sync.get("noteMarkdown", "")
            if note_md:
                path = ticker_dir / f"{ticker}_Note_{month}.md"
                path.write_text(note_md, encoding="utf-8")
                log.info(f"  Synced: {path.name}")
            sources_md = sync.get("sourcesMarkdown", "")
            if sources_md:
                path = ticker_dir / f"{ticker}_Sources_{month}.md"
                path.write_text(sources_md, encoding="utf-8")
                log.info(f"  Synced: {path.name}")
            changelog_md = sync.get("changelogMarkdown", "")
            if changelog_md:
                path = ticker_dir / f"{ticker}_Changelog.md"
                path.write_text(changelog_md, encoding="utf-8")
                log.info(f"  Synced: {path.name}")
            # Regenerate .docx from edited markdown
            if note_md:
                try:
                    company = ticker
                    try:
                        r = requests.get(f"{CHARLIE_API}/api/analysis/{ticker}", headers=_agent_headers(), timeout=10)
                        if r.status_code == 200:
                            company = r.json().get("company", ticker) or ticker
                    except Exception:
                        pass
                    # Find existing chart PNGs in the folder
                    chart_paths = [f for f in sorted(ticker_dir.glob(f"{ticker}_*_Breakdown.png"))]
                    docx_path = generate_note_docx(ticker, company, note_md, chart_paths, ticker_dir)
                    if docx_path:
                        log.info(f"  Synced: {docx_path.name}")
                except Exception as e:
                    log.warning(f"  DOCX regeneration failed: {e}")
            notify(f"*Charlie Agent:* Synced edited {ticker} note to iCloud (md + docx)")
    except Exception as e:
        log.debug(f"Sync check failed: {e}")


def import_existing_notes() -> None:
    """Scan all ticker folders for existing _Note_ .md files and upload to Charlie.

    Only imports notes that don't already exist in Charlie's database.
    Runs once on agent startup.
    """
    if not STOCKS_DIR.exists():
        return
    imported = 0
    for ticker_dir in sorted(STOCKS_DIR.iterdir()):
        if not ticker_dir.is_dir() or ticker_dir.name.startswith('.') or ticker_dir.name.startswith('_'):
            continue
        ticker = ticker_dir.name.upper()

        # Check if Charlie already has a note for this ticker
        try:
            r = requests.get(f"{CHARLIE_API}/api/notes/{ticker}", headers=_agent_headers(), timeout=10)
            if r.status_code == 200:
                continue  # Already has a note, skip
        except Exception:
            continue

        # Find _Note_ .md file
        note_files = sorted(ticker_dir.glob(f"{ticker}_Note_*.md"))
        if not note_files:
            continue
        note_file = note_files[-1]  # latest

        # Find matching sources and changelog
        sources_files = sorted(ticker_dir.glob(f"{ticker}_Sources_*.md"))
        changelog_files = sorted(ticker_dir.glob(f"{ticker}_Changelog*.md"))

        try:
            note_md = note_file.read_text(encoding="utf-8")
            sources_md = sources_files[-1].read_text(encoding="utf-8") if sources_files else ""
            changelog_md = changelog_files[-1].read_text(encoding="utf-8") if changelog_files else ""

            # Find chart PNGs
            charts_payload = []
            for cp in sorted(ticker_dir.glob(f"{ticker}_*_Breakdown.png")):
                chart_type = "revenue" if "Revenue" in cp.name else "profit"
                charts_payload.append({
                    "type": chart_type,
                    "filename": cp.name,
                    "data": base64.b64encode(cp.read_bytes()).decode("ascii"),
                })

            # Find docx
            docx_b64 = ""
            docx_files = sorted(ticker_dir.glob(f"{ticker}_Note_*.docx"))
            if docx_files:
                docx_b64 = base64.b64encode(docx_files[-1].read_bytes()).decode("ascii")

            # Extract version from filename or changelog
            version = "1.0"
            if changelog_md:
                import re as _re
                ver_match = _re.search(r'Version\s+(\d+\.\d+)', changelog_md)
                if ver_match:
                    version = ver_match.group(1)

            # Upload to Charlie
            resp = requests.post(
                f"{CHARLIE_API}/api/agent/save-note",
                json={
                    "noteId": str(uuid.uuid4()),
                    "ticker": ticker,
                    "version": version,
                    "noteMarkdown": note_md,
                    "sourcesMarkdown": sources_md,
                    "changelogMarkdown": changelog_md,
                    "noteDocx": docx_b64,
                    "charts": [{"type": c["type"], "filename": c["filename"]} for c in charts_payload],
                    "chartsData": charts_payload,
                    "metadata": {"importedFromICloud": True, "sourceFile": note_file.name},
                },
                headers=_agent_headers(),
                timeout=120,
            )
            if resp.status_code == 200:
                imported += 1
                log.info(f"  Imported existing note: {ticker} (v{version}) from {note_file.name}")
        except Exception as e:
            log.warning(f"  Failed to import {ticker} note: {e}")

    if imported:
        log.info(f"Imported {imported} existing notes to Charlie")


def import_existing_documents() -> None:
    """Bulk upload source documents (PDFs, xlsx) from iCloud folders to Charlie.

    Only uploads files not already in Charlie's document_files table.
    Scans main folder + Processed/ folder for each ticker.
    Runs once on agent startup.
    """
    if not STOCKS_DIR.exists():
        return

    UPLOAD_EXTS = {".pdf", ".xlsx", ".xls"}
    uploaded = 0
    skipped = 0

    for ticker_dir in sorted(STOCKS_DIR.iterdir()):
        if not ticker_dir.is_dir() or ticker_dir.name.startswith(".") or ticker_dir.name.startswith("_"):
            continue
        ticker = ticker_dir.name.upper()

        # Get list of documents already in Charlie for this ticker
        existing_filenames = set()
        try:
            r = requests.get(
                f"{CHARLIE_API}/api/documents/{ticker}",
                headers=_agent_headers(),
                timeout=15,
            )
            if r.status_code == 200:
                for doc in r.json().get("documents", []):
                    existing_filenames.add(doc.get("filename", ""))
        except Exception:
            pass

        # Scan main folder + Processed/ + any subfolders for source docs
        dirs_to_scan = [ticker_dir]
        for d in ticker_dir.iterdir():
            if d.is_dir() and d.name not in ("Prior Versions", ".icloud"):
                dirs_to_scan.append(d)

        docs_to_upload = []
        for scan_dir in dirs_to_scan:
            for f in sorted(scan_dir.iterdir()):
                if f.is_dir() or f.name.startswith(".") or f.name.startswith("~$"):
                    continue
                if f.suffix.lower() not in UPLOAD_EXTS:
                    continue
                if f.name in existing_filenames:
                    skipped += 1
                    continue
                if f.stat().st_size > 50 * 1024 * 1024:  # Skip >50MB
                    continue
                docs_to_upload.append(f)

        if not docs_to_upload:
            continue

        # Upload in batches of 3 to avoid overwhelming the backend
        for i in range(0, len(docs_to_upload), 3):
            batch = docs_to_upload[i : i + 3]
            documents = []
            for f in batch:
                try:
                    file_data = base64.b64encode(f.read_bytes()).decode("ascii")
                    ext = f.suffix.lower()
                    documents.append(
                        {
                            "filename": f.name,
                            "fileData": file_data,
                            "fileType": ext.lstrip("."),
                            "mimeType": {
                                ".pdf": "application/pdf",
                                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                ".xls": "application/vnd.ms-excel",
                            }.get(ext, "application/octet-stream"),
                        }
                    )
                except Exception as e:
                    log.warning(f"  Could not read {f.name}: {e}")

            if documents:
                try:
                    resp = requests.post(
                        f"{CHARLIE_API}/api/documents/save",
                        json={"ticker": ticker, "documents": documents},
                        headers=_agent_headers(),
                        timeout=120,
                    )
                    if resp.status_code == 200:
                        uploaded += len(documents)
                        for d in documents:
                            log.info(f"  Uploaded: {ticker}/{d['filename']}")
                except Exception as e:
                    log.warning(f"  Upload failed for {ticker}: {e}")

    if uploaded:
        log.info(f"Imported {uploaded} source documents to Charlie ({skipped} already existed)")


def check_and_fulfill_doc_requests() -> None:
    """Check backend for urgent document upload requests and fulfill them."""
    try:
        r = requests.get(
            f"{CHARLIE_API}/api/agent/doc-requests",
            headers=_agent_headers(),
            timeout=10,
        )
        if r.status_code != 200:
            return
        pending = r.json().get('requests', [])
        if not pending:
            return

        UPLOAD_EXTS = {".pdf", ".xlsx", ".xls", ".csv"}
        for req in pending:
            ticker = req['ticker']
            log.info(f"Urgent doc upload requested for {ticker}")
            ticker_dir = STOCKS_DIR / ticker
            if not ticker_dir.exists():
                log.warning(f"  Ticker folder not found: {ticker_dir}")
                continue

            # Get existing docs in Charlie to avoid duplicates
            existing_filenames = set()
            try:
                resp = requests.get(
                    f"{CHARLIE_API}/api/documents/{ticker}",
                    headers=_agent_headers(),
                    timeout=15,
                )
                if resp.status_code == 200:
                    for doc in resp.json().get("documents", []):
                        existing_filenames.add(doc.get("filename", ""))
            except Exception:
                pass

            log.info(f"  Found {ticker} folder, {len(existing_filenames)} already in DB, scanning...")
            # Scan main folder for source docs
            docs_to_upload = []
            for f in sorted(ticker_dir.iterdir()):
                if f.is_dir() or f.name.startswith('.') or f.name.startswith('~$'):
                    continue
                if f.suffix.lower() not in UPLOAD_EXTS:
                    continue
                if f.name in existing_filenames:
                    continue
                if f.stat().st_size > 50 * 1024 * 1024:
                    continue
                docs_to_upload.append(f)

            log.info(f"  {len(docs_to_upload)} docs to upload for {ticker}")
            uploaded = 0
            for i in range(0, len(docs_to_upload), 3):
                batch = docs_to_upload[i : i + 3]
                documents = []
                for f in batch:
                    try:
                        file_data = base64.b64encode(f.read_bytes()).decode("ascii")
                        ext = f.suffix.lower()
                        documents.append({
                            "filename": f.name,
                            "fileData": file_data,
                            "fileType": ext.lstrip("."),
                            "mimeType": {
                                ".pdf": "application/pdf",
                                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                ".xls": "application/vnd.ms-excel",
                                ".csv": "text/csv",
                            }.get(ext, "application/octet-stream"),
                        })
                    except Exception as e:
                        log.warning(f"  Could not read {f.name}: {e}")

                if documents:
                    try:
                        resp = requests.post(
                            f"{CHARLIE_API}/api/documents/save",
                            json={"ticker": ticker, "documents": documents},
                            headers=_agent_headers(),
                            timeout=120,
                        )
                        if resp.status_code == 200:
                            uploaded += len(documents)
                            for d in documents:
                                log.info(f"  Uploaded: {ticker}/{d['filename']}")
                    except Exception as e:
                        log.warning(f"  Upload failed for {ticker}: {e}")

            # Confirm upload complete
            try:
                requests.post(
                    f"{CHARLIE_API}/api/agent/doc-upload-complete",
                    json={"ticker": ticker, "count": uploaded},
                    headers=_agent_headers(),
                    timeout=10,
                )
            except Exception:
                pass
            log.info(f"  Uploaded {uploaded} docs for {ticker}")

    except Exception as e:
        log.error(f"Doc request check failed: {e}")
        import traceback
        traceback.print_exc()


def auto_unzip_stock_folders() -> None:
    """Find and extract any .zip files in stock ticker folders, then remove the zips."""
    if not STOCKS_DIR.exists():
        return
    for ticker_dir in sorted(STOCKS_DIR.iterdir()):
        if not ticker_dir.is_dir() or ticker_dir.name.startswith('.') or ticker_dir.name.startswith('_'):
            continue
        # Check main folder and immediate subdirectories for zips
        for search_dir in [ticker_dir] + [d for d in ticker_dir.iterdir() if d.is_dir() and d.name not in ('Processed', 'Prior Versions', '.icloud')]:
            for zf in sorted(search_dir.glob('*.zip')):
                if zf.name.startswith('.') or zf.name.startswith('~$'):
                    continue
                try:
                    # Wait for file to finish downloading (iCloud)
                    prev_size = -1
                    for _ in range(10):
                        cur_size = zf.stat().st_size
                        if cur_size == prev_size and cur_size > 0:
                            break
                        prev_size = cur_size
                        time.sleep(1)

                    with zipfile.ZipFile(zf, 'r') as z:
                        # Extract to the same directory as the zip
                        extracted = []
                        for member in z.namelist():
                            # Skip macOS metadata
                            if member.startswith('__MACOSX') or member.startswith('.'):
                                continue
                            z.extract(member, search_dir)
                            extracted.append(member)
                        log.info(f"  Unzipped {zf.name} -> {len(extracted)} files in {ticker_dir.name}/")
                    # Remove the zip after successful extraction
                    zf.unlink()
                    log.info(f"  Removed {zf.name}")
                except zipfile.BadZipFile:
                    log.warning(f"  Bad zip file, skipping: {zf.name}")
                except Exception as e:
                    log.warning(f"  Could not unzip {zf.name}: {e}")


def auto_unzip_catalyst_folders() -> None:
    """Find and extract any .zip files in catalyst topic folders, then remove the zips."""
    if not CATALYSTS_DIR.exists():
        return
    for ticker_dir in sorted(CATALYSTS_DIR.iterdir()):
        if not ticker_dir.is_dir() or ticker_dir.name.startswith('.'):
            continue
        # Check ticker folder and all topic subfolders
        for search_dir in [ticker_dir] + [d for d in ticker_dir.iterdir() if d.is_dir() and not d.name.startswith('.')]:
            for zf in sorted(search_dir.glob('*.zip')):
                if zf.name.startswith('.') or zf.name.startswith('~$'):
                    continue
                try:
                    # Wait for file to finish downloading (iCloud)
                    prev_size = -1
                    for _ in range(10):
                        cur_size = zf.stat().st_size
                        if cur_size == prev_size and cur_size > 0:
                            break
                        prev_size = cur_size
                        time.sleep(1)

                    with zipfile.ZipFile(zf, 'r') as z:
                        extracted = []
                        for member in z.namelist():
                            if member.startswith('__MACOSX') or member.startswith('.'):
                                continue
                            z.extract(member, search_dir)
                            extracted.append(member)
                        log.info(f"  Unzipped {zf.name} -> {len(extracted)} files in CATALYSTS/{ticker_dir.name}/{search_dir.name}/")
                    zf.unlink()
                    log.info(f"  Removed {zf.name}")
                except zipfile.BadZipFile:
                    log.warning(f"  Bad zip file, skipping: {zf.name}")
                except Exception as e:
                    log.warning(f"  Could not unzip {zf.name}: {e}")


def push_file_manifest() -> None:
    """Scan all ticker folders recursively and push file manifest to backend."""
    SCAN_EXTS = {'.pdf', '.xlsx', '.xls', '.csv', '.txt', '.md', '.docx', '.pptx', '.png'}
    SKIP_DIRS = {'.icloud'}
    manifest = {}
    for ticker_dir in sorted(STOCKS_DIR.iterdir()):
        if not ticker_dir.is_dir() or ticker_dir.name.startswith('.') or ticker_dir.name.startswith('_'):
            continue
        ticker = ticker_dir.name.upper()
        files = []

        # Recursively scan all files in the ticker folder
        for f in sorted(ticker_dir.rglob('*')):
            if f.is_dir():
                continue
            if f.name.startswith('.') or f.name.startswith('~$'):
                continue
            # Skip files inside .icloud dirs
            if any(part in SKIP_DIRS for part in f.parts):
                continue
            ext = f.suffix.lower()
            if ext not in SCAN_EXTS:
                continue
            # Determine folder label from relative path
            rel = f.relative_to(ticker_dir)
            if len(rel.parts) == 1:
                folder = 'main'
            else:
                folder = str(rel.parent)  # e.g., "Processed", "Prior Versions", "20260320 - Tony Lee..."
            files.append({
                'filename': f.name,
                'folder': folder,
                'path': str(rel),
                'size': f.stat().st_size,
                'extension': ext,
                'modified': datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
            })

        if files:
            manifest[ticker] = files

    try:
        requests.post(
            f"{CHARLIE_API}/api/agent/file-manifest",
            json={'manifest': manifest, 'timestamp': datetime.now().isoformat()},
            headers=_agent_headers(),
            timeout=15,
        )
    except Exception as e:
        log.debug(f"Manifest push failed: {e}")


# ---------------------------------------------------------------------------
# File reading
# ---------------------------------------------------------------------------


def read_ticker_files(
    ticker: str,
    include_processed: bool = False,
    file_selection: Optional[list[dict]] = None,
) -> list[dict]:
    """Read files from the ticker's iCloud folder.

    Args:
        ticker: Stock ticker
        include_processed: If True, also read from Processed/ subfolder
        file_selection: If provided, only include files matching this list.
                       Each entry: {'filename': 'name.pdf', 'folder': 'main'|'processed'}
    """
    ticker_dir = STOCKS_DIR / ticker
    if not ticker_dir.exists():
        raise FileNotFoundError(f"No iCloud folder found for {ticker} at {ticker_dir}")

    MIME_MAP = {
        ".pdf": "application/pdf",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    }
    ALLOWED_EXTS = set(MIME_MAP.keys())

    # Build a set of (filename, folder) for fast lookup when file_selection is provided
    selection_set: Optional[set[tuple[str, str]]] = None
    if file_selection:
        selection_set = {(entry['filename'], entry.get('folder', 'main')) for entry in file_selection}

    def _should_include(filename: str, folder: str) -> bool:
        if selection_set is not None:
            return (filename, folder) in selection_set
        return True

    def _read_dir(directory: Path, folder_label: str) -> list[dict]:
        result: list[dict] = []
        if not directory.exists():
            return result
        for f in sorted(directory.iterdir()):
            if f.is_dir():
                continue
            if f.name.startswith(".") or f.name.startswith("~$"):
                continue
            ext = f.suffix.lower()
            if ext not in ALLOWED_EXTS:
                continue
            if not _should_include(f.name, folder_label):
                continue
            try:
                raw = f.read_bytes()
                result.append(
                    {
                        "filename": f.name,
                        "path": str(f),
                        "folder": folder_label,
                        "extension": ext,
                        "size": len(raw),
                        "data": base64.b64encode(raw).decode("ascii"),
                        "mime_type": MIME_MAP.get(ext, "application/octet-stream"),
                    }
                )
                log.debug(f"  Read {f.name} ({len(raw):,} bytes) [{folder_label}]")
            except Exception as e:
                log.warning(f"  Could not read {f.name}: {e}")
        return result

    files: list[dict] = []

    # Main folder (always scan)
    files.extend(_read_dir(ticker_dir, "main"))

    # Determine which subfolders to scan
    subfolders_to_scan: set[str] = set()
    if include_processed:
        subfolders_to_scan.add("Processed")
    # If file_selection specifies files from specific folders, scan those too
    if selection_set:
        for _, folder in selection_set:
            if folder != "main":
                subfolders_to_scan.add(folder)

    # Scan each required subfolder
    for subfolder in subfolders_to_scan:
        sub_dir = ticker_dir / subfolder
        if sub_dir.exists() and sub_dir.is_dir():
            files.extend(_read_dir(sub_dir, subfolder))

    return files


# ---------------------------------------------------------------------------
# Excel text extraction
# ---------------------------------------------------------------------------


def extract_excel_text(filepath: str) -> Optional[str]:
    """Extract text content from Excel files. Uses zipfile+XML as fallback."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
        text_parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"\n=== Sheet: {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join(str(c) if c is not None else "" for c in row)
                if row_text.strip():
                    text_parts.append(row_text)
                    row_count += 1
                if row_count > 2000:
                    text_parts.append("[... truncated at 2000 rows ...]")
                    break
        wb.close()
        return "\n".join(text_parts)
    except Exception as e:
        log.debug(f"openpyxl failed for {filepath}: {e}, trying zipfile fallback")
        try:
            return _extract_xlsx_via_zip(filepath)
        except Exception as e2:
            log.warning(f"Could not extract Excel text from {filepath}: {e2}")
            return None


def _extract_xlsx_via_zip(filepath: str) -> Optional[str]:
    """Parse xlsx as zip + XML when openpyxl fails."""
    ns = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"

    with zipfile.ZipFile(filepath) as z:
        # Read shared strings
        shared: dict[int, str] = {}
        try:
            tree = ElementTree.parse(z.open("xl/sharedStrings.xml"))
            for i, si in enumerate(tree.findall(f".//{ns}si")):
                texts = si.findall(f".//{ns}t")
                shared[i] = "".join(t.text or "" for t in texts)
        except (KeyError, ElementTree.ParseError):
            pass

        text_parts: list[str] = []
        for name in sorted(z.namelist()):
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"):
                sheet_tree = ElementTree.parse(z.open(name))
                rows = sheet_tree.findall(f".//{ns}row")
                for row in rows:
                    cells: list[str] = []
                    for cell in row.findall(f"{ns}c"):
                        val = cell.find(f"{ns}v")
                        cell_type = cell.get("t", "")
                        if val is not None and val.text:
                            if cell_type == "s":
                                cells.append(shared.get(int(val.text), val.text))
                            else:
                                cells.append(val.text)
                        else:
                            cells.append("")
                    row_text = "\t".join(cells)
                    if row_text.strip():
                        text_parts.append(row_text)

        return "\n".join(text_parts) if text_parts else None


# ---------------------------------------------------------------------------
# Claude API call
# ---------------------------------------------------------------------------


MAX_TOKENS_PER_BATCH = 170_000  # Leave room for prompt + response under 200K limit
APPROX_TOKENS_PER_BYTE = 0.4  # rough estimate for PDF base64 -> tokens


def _estimate_tokens(files: list[dict]) -> int:
    """Rough estimate of token count for a set of files."""
    total = 0
    for f in files:
        if f["extension"] == ".pdf":
            # PDF base64 is ~1.33x raw size, tokens ~ 0.4 per byte of content
            total += int(f["size"] * APPROX_TOKENS_PER_BYTE)
        else:
            total += int(len(f.get("data", "")) * 0.3)
    return total


def _split_into_batches(files: list[dict]) -> list[list[dict]]:
    """Split files into batches that fit within Claude's context window."""
    batches: list[list[dict]] = []
    current_batch: list[dict] = []
    current_tokens = 0

    for f in files:
        est = _estimate_tokens([f])
        if current_tokens + est > MAX_TOKENS_PER_BATCH and current_batch:
            batches.append(current_batch)
            current_batch = []
            current_tokens = 0
        current_batch.append(f)
        current_tokens += est

    if current_batch:
        batches.append(current_batch)
    return batches if batches else [files]


def call_claude(
    files: list[dict],
    ticker: str,
    company: str,
    sector: str,
    existing_note: Optional[str] = None,
    mode: str = "new",
    api_key: Optional[str] = None,
    on_progress: Optional[callable] = None,
) -> str:
    """Call Claude API with source files, auto-batching if too large."""
    est_tokens = _estimate_tokens(files)
    batches = _split_into_batches(files)

    if len(batches) > 1:
        log.info(f"  Documents exceed context limit (~{est_tokens:,} tokens). Splitting into {len(batches)} batches.")

    # Process first batch — generates the full note
    result_text = _call_claude_single(
        batches[0], ticker, company, sector, existing_note, mode, api_key
    )

    # Process subsequent batches — merge additional data into existing note
    for i, batch in enumerate(batches[1:], 2):
        log.info(f"  Processing batch {i}/{len(batches)}...")
        if on_progress:
            on_progress(f"Processing batch {i}/{len(batches)}")

        # Extract the note from previous result to use as context
        m = re.search(r'===NOTE_START===\s*(.*?)\s*===NOTE_END===', result_text, re.DOTALL)
        prev_note = m.group(1).strip() if m else result_text[:8000]

        result_text = _call_claude_single(
            batch, ticker, company, sector,
            existing_note=prev_note,
            mode="update",
            api_key=api_key,
        )

    return result_text


def _call_claude_single(
    files: list[dict],
    ticker: str,
    company: str,
    sector: str,
    existing_note: Optional[str] = None,
    mode: str = "new",
    api_key: Optional[str] = None,
) -> str:
    """Call Claude API with a single batch of files."""
    client = anthropic.Anthropic(api_key=api_key)

    content: list[dict] = []
    pdf_count = 0

    # Attach documents
    for f in files:
        if f["extension"] == ".pdf":
            doc_block = {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": f["data"],
                },
            }
            # Claude allows max 4 blocks with cache_control
            if pdf_count < 4:
                doc_block["cache_control"] = {"type": "ephemeral"}
            pdf_count += 1
            content.append(doc_block)
            content.append({"type": "text", "text": f"[Document: {f['filename']}]"})

        elif f["extension"] in (".xlsx", ".xls"):
            text = extract_excel_text(f["path"])
            if text:
                # Truncate very large spreadsheets
                if len(text) > 80_000:
                    text = text[:80_000] + "\n[... truncated ...]"
                content.append(
                    {"type": "text", "text": f"[Excel Model: {f['filename']}]\n{text}"}
                )

        elif f["extension"] in (".csv", ".txt", ".md"):
            try:
                text = base64.b64decode(f["data"]).decode("utf-8", errors="replace")
                if len(text) > 80_000:
                    text = text[:80_000] + "\n[... truncated ...]"
                content.append({"type": "text", "text": f"[{f['filename']}]\n{text}"})
            except Exception:
                pass

        elif f["extension"] == ".docx":
            # Try to extract text from docx
            docx_text = _extract_docx_text(f["path"])
            if docx_text:
                if len(docx_text) > 80_000:
                    docx_text = docx_text[:80_000] + "\n[... truncated ...]"
                content.append(
                    {"type": "text", "text": f"[Word Doc: {f['filename']}]\n{docx_text}"}
                )

    # Build the existing-thesis context if we have analysis data from Charlie
    existing_thesis_context = ""
    try:
        resp = requests.get(
            f"{CHARLIE_API}/api/analysis/{ticker}",
            headers=_agent_headers(),
            timeout=15,
        )
        if resp.status_code == 200:
            analysis = resp.json()
            thesis = analysis.get("thesis", {})
            signposts = analysis.get("signposts", [])
            threats = analysis.get("threats", [])
            conclusion = analysis.get("conclusion", "")
            if thesis or signposts or threats:
                existing_thesis_context = f"""
EXISTING STRUCTURED THESIS (use as foundation, expand with document details):
Summary: {thesis.get('summary', '')}
Pillars: {json.dumps(thesis.get('pillars', []), indent=2)}
Signposts: {json.dumps(signposts, indent=2)}
Threats: {json.dumps(threats, indent=2)}
Conclusion: {conclusion}
"""
    except Exception:
        pass

    update_context = ""
    if existing_note and mode == "update":
        update_context = f"""
EXISTING NOTE (update this, don't rewrite from scratch):
{existing_note[:8000]}

Update the note with new information from the source documents. Add a "What's Changed" section at the top. Update any data points, estimates, or catalysts that have changed.
"""

    prompt = f"""You are a senior equity research analyst writing a comprehensive investment research note.

TICKER: {ticker}
COMPANY: {company}
SECTOR: {sector}

{RESEARCH_NOTE_PLAYBOOK}

{existing_thesis_context}

{update_context}

Using the source documents provided, write a complete equity research note in markdown format.

The note should be 8-12 pages when printed, with these sections:
1. Executive Summary / Investment Thesis
2. Business Overview & Segment Breakdown
3. Key Revenue & Earnings Drivers
4. What the Street Is Debating
5. Catalyst Calendar
6. Valuation Context
7. Risks
8. Bottom Line

Also provide:
- Revenue segment data for donut chart (JSON array: [{{"segment": "name", "revenue": number_in_millions}}])
- Profit/Operating Income segment data for donut chart (JSON array: [{{"segment": "name", "profit": number_in_millions}}])
  IMPORTANT: Revenue and profit MUST be different numbers. Profit means operating income, EBIT, NOI, or segment profit — NOT revenue.
  If segment-level profit data is not available in the source documents, return an empty array [] instead of reusing revenue numbers.
  For REITs, use NOI by segment. For industrials, use segment operating profit. For pharma, use segment operating income.

IMPORTANT:
- Source attributions go in a SEPARATE sources document, NOT in the main note
- In the sources doc, list section-by-section which reports informed each claim (broker name + date + page ref)
- In the main note, NEVER mention broker names, analyst names, or firm-specific targets
- Do NOT use Street opinion as thesis support (e.g. "the Street is constructive", "consensus targets suggest upside")
- OK in Valuation section: factual framing like "applying 11-12x to normalized EPS implies $X" (your own math)
- Nudge displayed numbers +/- $200-300M from model data to avoid exact replication

Return your response in this exact format:

===NOTE_START===
[full markdown note here]
===NOTE_END===

===SOURCES_START===
[sources document: section-by-section, which reports informed each claim, broker name + date + page ref]
===SOURCES_END===

===REVENUE_CHART_DATA===
[JSON array of revenue segments]
===REVENUE_CHART_END===

===PROFIT_CHART_DATA===
[JSON array of profit segments]
===PROFIT_CHART_END===
"""

    content.append({"type": "text", "text": prompt})

    log.info(f"Calling Claude with {len(content)} content blocks...")

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=16384,
        system="You are a senior equity research analyst. Write thorough, data-driven research notes. Be precise with numbers. No sellside attribution in the main note.",
        messages=[{"role": "user", "content": content}],
    )

    return response.content[0].text


def _extract_docx_text(filepath: str) -> Optional[str]:
    """Extract plain text from a .docx file."""
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        with zipfile.ZipFile(filepath) as z:
            tree = ElementTree.parse(z.open("word/document.xml"))
            paragraphs = tree.findall(f".//{ns}p")
            texts: list[str] = []
            for p in paragraphs:
                runs = p.findall(f".//{ns}t")
                para_text = "".join(r.text or "" for r in runs)
                if para_text.strip():
                    texts.append(para_text)
            return "\n".join(texts) if texts else None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Response parsing
# ---------------------------------------------------------------------------


def parse_response(text: str) -> tuple[str, str, list[dict], list[dict]]:
    """Parse the LLM response into (note_md, sources_md, revenue_data, profit_data)."""
    note_md = ""
    sources_md = ""
    revenue_data: list[dict] = []
    profit_data: list[dict] = []

    m = re.search(r"===NOTE_START===\s*(.*?)\s*===NOTE_END===", text, re.DOTALL)
    note_md = m.group(1).strip() if m else text

    m = re.search(r"===SOURCES_START===\s*(.*?)\s*===SOURCES_END===", text, re.DOTALL)
    sources_md = m.group(1).strip() if m else ""

    m = re.search(
        r"===REVENUE_CHART_DATA===\s*(.*?)\s*===REVENUE_CHART_END===", text, re.DOTALL
    )
    if m:
        try:
            revenue_data = json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            log.warning("Could not parse revenue chart data JSON")

    m = re.search(
        r"===PROFIT_CHART_DATA===\s*(.*?)\s*===PROFIT_CHART_END===", text, re.DOTALL
    )
    if m:
        try:
            profit_data = json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            log.warning("Could not parse profit chart data JSON")

    return note_md, sources_md, revenue_data, profit_data


# ---------------------------------------------------------------------------
# Chart generation (donut charts matching backend format)
# ---------------------------------------------------------------------------


def generate_donut_chart(
    ticker: str,
    chart_type: str,
    data: list[dict],
    value_key: str,
    output_dir: Path,
) -> Optional[Path]:
    """Generate a donut chart PNG in the same style as the backend."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.patheffects as pe
    import matplotlib.pyplot as plt
    import numpy as np

    # Filter out segments with negative or zero values (e.g., operating losses)
    filtered = [(d.get("segment", ""), d.get(value_key, d.get("revenue", d.get("profit", 0)))) for d in data]
    filtered = [(label, val) for label, val in filtered if val and val > 0]
    if not filtered:
        return None
    labels = [f[0] for f in filtered]
    values = [f[1] for f in filtered]

    total = sum(values)
    colors = [
        "#5DADE2",
        "#F7DC6F",
        "#F1948A",
        "#7DCEA0",
        "#BB8FCE",
        "#85C1E9",
        "#F8C471",
        "#82E0AA",
    ]

    fig, ax = plt.subplots(figsize=(10, 8), facecolor="white")
    wedges, _ = ax.pie(
        values,
        labels=None,
        colors=colors[: len(values)],
        startangle=90,
        wedgeprops={"width": 0.58, "edgecolor": "none", "linewidth": 0},
    )

    # Center circle and text
    centre_circle = plt.Circle((0, 0), 0.30, fc="white")
    ax.add_artist(centre_circle)
    total_str = f"\\${total/1000:.1f}B" if total >= 1000 else f"\\${total:.0f}M"
    ax.text(
        0, 0.05, "Total", ha="center", va="center", fontsize=12, color="#333333", fontweight="normal"
    )
    ax.text(
        0, -0.08, total_str, ha="center", va="center", fontsize=16, color="#333333", fontweight="bold"
    )

    # Inside labels (value + percentage on each wedge)
    for i, (wedge, value) in enumerate(zip(wedges, values)):
        pct = value / total * 100
        ang = (wedge.theta2 - wedge.theta1) / 2.0 + wedge.theta1
        x = 0.70 * np.cos(np.deg2rad(ang))
        y = 0.70 * np.sin(np.deg2rad(ang))
        val_str = f"\\${value/1000:.1f}B" if value >= 1000 else f"\\${value:.0f}M"
        fontsize = 11 if pct > 15 else 10 if pct > 8 else 9
        if pct >= 3:
            ax.text(
                x,
                y,
                f"{val_str}\n({pct:.1f}%)",
                ha="center",
                va="center",
                fontsize=fontsize,
                fontweight="bold",
                color="white",
                path_effects=[pe.withStroke(linewidth=2, foreground="black")],
            )

    # Outside labels (segment names)
    for i, (wedge, label) in enumerate(zip(wedges, labels)):
        ang = (wedge.theta2 - wedge.theta1) / 2.0 + wedge.theta1
        x = 1.15 * np.cos(np.deg2rad(ang))
        y = 1.15 * np.sin(np.deg2rad(ang))
        ax.text(
            x, y, label, ha="center", va="center", fontsize=10, fontweight="bold", color="#333333"
        )

    ax.set_title(
        f"{ticker} -- {chart_type} Breakdown",
        fontsize=14,
        fontweight="bold",
        color="#333333",
        pad=20,
    )
    ax.set_aspect("equal")
    plt.tight_layout()

    filename = f"{ticker}_{chart_type.replace(' ', '_')}_Breakdown.png"
    output_path = output_dir / filename
    plt.savefig(output_path, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    log.info(f"  Chart saved: {filename}")
    return output_path


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------


def _add_md_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, converting **bold** and *italic* markdown to Word formatting."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def generate_note_docx(
    ticker: str,
    company: str,
    markdown_text: str,
    chart_paths: list[Path],
    output_dir: Path,
) -> Optional[Path]:
    """Generate Word document from markdown with embedded charts.

    Uses Calibri 11pt, all black text (matching Tony's md_to_docx_v2.py style).
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        log.error("python-docx not installed -- skipping .docx generation")
        return None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Title
    title = doc.add_heading(f"{ticker} -- {company}", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Date line
    doc.add_paragraph(f"Equity Research Note -- {datetime.now().strftime('%B %Y')}")

    # Parse and render markdown
    lines = markdown_text.split("\n")
    charts_inserted = False
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip().replace('**', ''), level=3)
            h.paragraph_format.keep_with_next = True
            for r in h.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith("## "):
            section_title = line[3:].strip().replace('**', '')
            h = doc.add_heading(section_title, level=2)
            h.paragraph_format.keep_with_next = True
            for r in h.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
            # Insert charts after Business Overview section
            if not charts_inserted and (
                "business" in section_title.lower() or "segment" in section_title.lower()
            ):
                charts_inserted = True
                # We'll insert charts after this section's content ends
                # (at the next heading or after a blank line gap)

        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.paragraph_format.keep_with_next = True
            for r in h.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_md_runs(p, line[2:].strip())

        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_md_runs(p, re.sub(r"^\d+\.\s", "", line).strip())

        elif line.startswith("|") and "|" in line[1:]:
            # Simple markdown table handling
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[\s\-:|]+\|$", lines[i]):  # skip separator rows
                    table_lines.append(lines[i])
                i += 1
            i -= 1  # will be incremented below

            if table_lines:
                # Parse cells
                parsed_rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip("|").split("|")]
                    parsed_rows.append(cells)

                if parsed_rows:
                    num_cols = max(len(r) for r in parsed_rows)
                    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                    table.style = "Table Grid"
                    table.autofit = False

                    # Distribute column widths across page (6.5in usable)
                    # Give last column extra space for text-heavy content
                    total_width = 6.5
                    if num_cols <= 1:
                        col_widths = [total_width]
                    elif num_cols <= 3:
                        col_widths = [total_width / num_cols] * num_cols
                    else:
                        # Last column gets 35% of width; rest split evenly
                        last_w = total_width * 0.35
                        even_w = (total_width - last_w) / (num_cols - 1)
                        col_widths = [even_w] * (num_cols - 1) + [last_w]

                    for ci in range(num_cols):
                        table.columns[ci].width = Inches(col_widths[ci])

                    for ri, row_cells in enumerate(parsed_rows):
                        for ci, cell_text in enumerate(row_cells):
                            if ci < num_cols:
                                cell = table.cell(ri, ci)
                                cell.width = Inches(col_widths[ci])
                                cell.text = cell_text.replace('**', '').replace('*', '')
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(10)
                                        r.font.name = "Calibri"
                                        r.font.color.rgb = RGBColor(0, 0, 0)
                    # Bold header row with dark background
                    if len(parsed_rows) > 0:
                        for ci in range(num_cols):
                            hdr_cell = table.cell(0, ci)
                            # Set header background to #2C3E50
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), '2C3E50')
                            shading.set(qn('w:val'), 'clear')
                            hdr_cell._tc.get_or_add_tcPr().append(shading)
                            for p in hdr_cell.paragraphs:
                                for r in p.runs:
                                    r.bold = True
                                    r.font.color.rgb = RGBColor(255, 255, 255)

        elif line.strip():
            p = doc.add_paragraph()
            _add_md_runs(p, line)

        i += 1

    # Embed charts at the end (or after Business Overview if we found it)
    if chart_paths:
        ch = doc.add_heading("Charts", level=2)
        ch.paragraph_format.keep_with_next = True
        for cp in chart_paths:
            if cp and cp.exists():
                try:
                    # Add chart title that stays with the image
                    cap = doc.add_paragraph(cp.stem.replace("_", " "))
                    cap.paragraph_format.keep_with_next = True
                    doc.add_picture(str(cp), width=Inches(6))
                    doc.add_paragraph("")  # spacing
                except Exception as e:
                    log.warning(f"Could not embed chart {cp.name}: {e}")

    month = datetime.now().strftime("%b%Y")
    filename = f"{ticker}_Note_{month}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    log.info(f"  DOCX saved: {filename}")
    return output_path


# ---------------------------------------------------------------------------
# File organization (Prior Versions, Processed)
# ---------------------------------------------------------------------------


def organize_files(ticker_dir: Path, source_files: list[dict], skip_move: bool = False) -> None:
    """Archive existing deliverables to Prior Versions/, move sources to Processed/.

    Follows Tony's memory preferences:
    - Prior versions get date suffix
    - Main folder should only contain latest deliverables
    - Source files go to Processed/

    Args:
        ticker_dir: Path to the ticker's iCloud folder
        source_files: List of file dicts that were read as sources
        skip_move: If True, archive prior deliverables but do NOT move source
                   files to Processed/ (used for reprocess mode where files
                   are already in Processed/)
    """
    processed_dir = ticker_dir / "Processed"
    processed_dir.mkdir(exist_ok=True)

    prior_dir = ticker_dir / "Prior Versions"
    prior_dir.mkdir(exist_ok=True)

    date_suffix = datetime.now().strftime("%Y%m%d")

    # 1. Archive existing deliverables (notes, sources, changelogs, charts)
    for f in sorted(ticker_dir.iterdir()):
        if f.is_dir():
            continue
        name_lower = f.name.lower()
        ext = f.suffix.lower()

        is_deliverable = False
        # Note / Sources / Changelog markdown or docx
        if ext in (".md", ".docx") and any(
            kw in f.name for kw in ("_Note_", "_Sources_", "_Changelog")
        ):
            is_deliverable = True
        # Chart PNGs
        if ext == ".png" and any(
            kw in f.name for kw in ("Breakdown", "Revenue", "Profit")
        ):
            is_deliverable = True

        if is_deliverable:
            dest = prior_dir / f"{f.stem}_{date_suffix}{f.suffix}"
            # Avoid overwriting if archived today already
            if dest.exists():
                dest = prior_dir / f"{f.stem}_{date_suffix}_{uuid.uuid4().hex[:6]}{f.suffix}"
            shutil.move(str(f), str(dest))
            log.info(f"  Archived: {f.name} -> Prior Versions/")

    # 2. Move source files to Processed/ (unless skip_move is set)
    if not skip_move:
        for sf in source_files:
            src = Path(sf["path"])
            if src.exists() and src.parent == ticker_dir:
                dest = processed_dir / src.name
                if dest.exists():
                    dest = processed_dir / f"{src.stem}_{date_suffix}{src.suffix}"
                shutil.move(str(src), str(dest))
                log.info(f"  Processed: {src.name}")


# ---------------------------------------------------------------------------
# Incremental update support
# ---------------------------------------------------------------------------


def find_existing_note(ticker_dir: Path, ticker: str) -> Optional[str]:
    """Find and read the most recent note markdown in the ticker folder.

    Used for incremental updates -- if a prior note exists, we pass it to
    Claude so it can update rather than rewrite from scratch.
    """
    candidates: list[Path] = []
    for f in ticker_dir.iterdir():
        if f.is_dir():
            continue
        if f.suffix.lower() == ".md" and "_Note_" in f.name:
            candidates.append(f)

    if not candidates:
        return None

    # Pick the most recently modified
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    try:
        return candidates[0].read_text(encoding="utf-8")
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Version management
# ---------------------------------------------------------------------------


def determine_version(ticker_dir: Path, ticker: str, mode: str) -> str:
    """Determine the next version number based on existing notes."""
    if mode == "new":
        return "1.0"

    # Look for existing changelog to find latest version
    for f in ticker_dir.iterdir():
        if f.is_dir():
            continue
        if "_Changelog" in f.name and f.suffix.lower() == ".md":
            try:
                text = f.read_text(encoding="utf-8")
                versions = re.findall(r"## Version (\d+\.\d+)", text)
                if versions:
                    latest = versions[0]
                    parts = latest.split(".")
                    parts[-1] = str(int(parts[-1]) + 1)
                    return ".".join(parts)
            except Exception:
                pass

    return "1.0"


# ---------------------------------------------------------------------------
# Job orchestrator
# ---------------------------------------------------------------------------


def process_note_job(job: dict, api_key: str) -> None:
    """Process a single note generation job end-to-end."""
    job_id = job.get("id", str(uuid.uuid4()))
    ticker = job.get("ticker", "").upper()
    mode = "new"
    steps_detail = job.get("steps_detail", {})
    if isinstance(steps_detail, str):
        try:
            steps_detail = json.loads(steps_detail)
        except Exception:
            steps_detail = {}
    mode = steps_detail.get("mode", "new")
    file_selection = steps_detail.get("fileSelection", [])  # list of {filename, folder}
    reprocess = steps_detail.get("reprocess", False)

    log.info(f"=== Processing {ticker} (job {job_id[:12]}..., mode={mode}, reprocess={reprocess}) ===")
    if file_selection:
        log.info(f"  File selection: {len(file_selection)} files specified")
    notify(f"*Charlie Agent:* Picked up {ticker} note job ({mode})")

    try:
        # Step 1: Claim and read files
        update_job_progress(job_id, "running", "Reading local files", 5)

        ticker_dir = STOCKS_DIR / ticker
        if not ticker_dir.exists():
            raise FileNotFoundError(f"No iCloud folder for {ticker} at {ticker_dir}")

        files = read_ticker_files(
            ticker,
            include_processed=reprocess or bool(file_selection),
            file_selection=file_selection if file_selection else None,
        )
        if not files:
            raise RuntimeError(f"No source files found in {ticker_dir}")

        log.info(f"  Found {len(files)} source files")
        update_job_progress(job_id, "running", f"Found {len(files)} files, calling Claude", 15)

        # Get company name from Charlie
        company = ticker
        try:
            resp = requests.get(
                f"{CHARLIE_API}/api/analysis/{ticker}",
                headers=_agent_headers(),
                timeout=15,
            )
            if resp.status_code == 200:
                company = resp.json().get("company", ticker) or ticker
        except Exception:
            pass

        sector = get_sector(ticker)

        # Check for existing note (for incremental updates)
        existing_note = None
        if mode == "update":
            existing_note = find_existing_note(ticker_dir, ticker)
            if existing_note:
                log.info("  Found existing note for incremental update")

        # Step 2: Call Claude (auto-batches if too many documents)
        update_job_progress(job_id, "running", "Generating research note via Claude", 20)
        response_text = call_claude(
            files=files,
            ticker=ticker,
            company=company,
            sector=sector,
            existing_note=existing_note,
            mode=mode,
            api_key=api_key,
            on_progress=lambda msg: update_job_progress(job_id, "running", msg, 40),
        )

        # Step 3: Parse response
        update_job_progress(job_id, "running", "Parsing response", 60)
        note_md, sources_md, revenue_data, profit_data = parse_response(response_text)

        if not note_md or len(note_md) < 500:
            log.warning(f"  Note seems too short ({len(note_md)} chars), using raw response")
            if len(response_text) > len(note_md):
                note_md = response_text

        log.info(f"  Note: {len(note_md):,} chars, Sources: {len(sources_md):,} chars")
        log.info(
            f"  Revenue segments: {len(revenue_data)}, Profit segments: {len(profit_data)}"
        )

        # Step 4: Archive prior versions + move sources BEFORE creating ANY new files
        update_job_progress(job_id, "running", "Organizing files", 70)
        month = datetime.now().strftime("%b%Y")
        version = determine_version(ticker_dir, ticker, mode)
        organize_files(ticker_dir, files, skip_move=reprocess)

        # Step 5: Generate charts (AFTER archival so they don't get archived)
        update_job_progress(job_id, "running", "Generating charts", 75)
        chart_paths: list[Path] = []

        if revenue_data:
            cp = generate_donut_chart(ticker, "Revenue", revenue_data, "revenue", ticker_dir)
            if cp:
                chart_paths.append(cp)

        if profit_data:
            # Validate: skip if profit data is identical to revenue data (LLM reused same numbers)
            rev_values = sorted([d.get('revenue', d.get('profit', 0)) for d in revenue_data]) if revenue_data else []
            prof_values = sorted([d.get('profit', d.get('revenue', 0)) for d in profit_data])
            if rev_values and prof_values and rev_values == prof_values:
                log.warning("  Profit chart data identical to revenue — skipping profit chart")
            else:
                cp = generate_donut_chart(ticker, "Operating Profit", profit_data, "profit", ticker_dir)
                if cp:
                    chart_paths.append(cp)

        # Step 6: Generate DOCX with embedded charts
        update_job_progress(job_id, "running", "Building .docx", 80)
        docx_path = generate_note_docx(ticker, company, note_md, chart_paths, ticker_dir)

        # Write new deliverables
        note_path = ticker_dir / f"{ticker}_Note_{month}.md"
        note_path.write_text(note_md, encoding="utf-8")
        log.info(f"  Saved: {note_path.name}")

        sources_path = ticker_dir / f"{ticker}_Sources_{month}.md"
        sources_path.write_text(sources_md, encoding="utf-8")
        log.info(f"  Saved: {sources_path.name}")

        # Changelog
        now_str = datetime.now().strftime("%Y-%m-%d")
        changelog_md = (
            f"# {ticker} Changelog\n\n"
            f"## Version {version} -- {now_str}\n"
            f"- {'Updated' if mode == 'update' else 'Initial'} research note generated\n"
            f"- {len(files)} source documents processed\n"
            f"- Generated via Charlie\n"
        )

        # Append old changelog if this is an update
        if mode == "update":
            prior_dir = ticker_dir / "Prior Versions"
            for f in sorted(prior_dir.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
                if "_Changelog" in f.name and f.suffix.lower() == ".md":
                    try:
                        old_text = f.read_text(encoding="utf-8")
                        # Strip the top-level heading and append
                        old_body = old_text.split("\n", 1)[1] if "\n" in old_text else old_text
                        changelog_md += "\n" + old_body
                    except Exception:
                        pass
                    break

        changelog_path = ticker_dir / f"{ticker}_Changelog.md"
        changelog_path.write_text(changelog_md, encoding="utf-8")
        log.info(f"  Saved: {changelog_path.name}")

        # Step 7: Upload to Charlie
        update_job_progress(job_id, "running", "Uploading to Charlie", 90)
        upload_note_to_charlie(
            ticker=ticker,
            note_md=note_md,
            sources_md=sources_md,
            changelog_md=changelog_md,
            docx_path=docx_path,
            chart_paths=chart_paths,
            version=version,
            metadata={
                "mode": mode,
                "documentsProcessed": len(files),
                "generatedLocally": True,
                "agentId": _agent_id(),
                "charCount": len(note_md),
            },
        )

        # Done
        update_job_progress(job_id, "complete", "Complete", 100)
        complete_job(job_id, {"ticker": ticker, "version": version, "charCount": len(note_md)})
        log.info(f"=== {ticker} complete (v{version}) ===\n")
        notify(f"*Charlie Agent:* {ticker} note complete (v{version})\n{len(files)} docs processed, {len(note_md):,} chars")

    except Exception as e:
        log.error(f"Job {job_id[:12]}... failed: {e}")
        import traceback

        traceback.print_exc()
        update_job_progress(job_id, "failed", "Failed", None, error=str(e))
        notify(f"*Charlie Agent:* {ticker} note FAILED\n{str(e)[:200]}")


# ---------------------------------------------------------------------------
# Catalyst synthesis
# ---------------------------------------------------------------------------

CATALYSTS_DIR = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/CATALYSTS"
SUMMARIES_DIR = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/SUMMARIES"

CATALYST_SYNTHESIS_PROMPT = """You are a senior equity research analyst writing a PROPRIETARY synthesis report for your portfolio manager. This must read as YOUR OWN original work product.

## ATTRIBUTION RULES — ABSOLUTE, NO EXCEPTIONS
- ZERO references to any broker, bank, sellside firm, or analyst by name. Not "Wolfe Research", not "Goldman Sachs", not "JPMorgan", not "Morgan Stanley", not ANY firm. If you write a firm name, you have FAILED.
- ZERO phrases like "Wolfe estimates", "according to [firm]", "[firm] notes", "[firm] identifies", "[firm] describes". Replace ALL such attributions with first-person voice: "I estimate", "my analysis shows", "I note".
- ZERO analyst counts, ratings, consensus targets, or price targets used as arguments.
- ZERO sellside sentiment used as thesis support ("the Street is constructive", "analysts favor").
- You may use "consensus estimates" or "street estimates" as a DATA LABEL in valuation math only.
- Write ENTIRELY in FIRST PERSON: "I think", "My view is", "I see the risk-reward as", "I estimate".
- Present all data, analysis, and conclusions as YOUR OWN proprietary work.

## SYNTHESIS RULES — USE ALL SOURCES EQUALLY
- You are receiving MULTIPLE source documents. Synthesize insights from ALL of them, not just the first or longest one.
- Weight each source proportionally. If 5 documents are provided, each should contribute meaningfully.
- When sources disagree, note the divergence and state YOUR view.
- Extract the BEST data points from EACH source and weave them into a unified narrative.

## TONE & STYLE
- Confident, direct analyst voice writing to their PM
- No AI filler phrases ("it's worth noting", "delve into", "poised to", "navigate the landscape")
- Lead with conclusions, support with evidence
- Be specific with numbers, data points, and dates

## LENGTH
{length_instruction}

## TOPIC
Ticker: {ticker}
Topic: {topic}
{custom_instructions}

## SOURCE DOCUMENTS
{source_content}

Write the synthesis report now. Start with a clear title line, then the analysis. Remember: ZERO firm names, ALL first person, synthesize ALL documents equally."""

CATALYST_LENGTH_PRESETS = {
    'quick': 'Write a single paragraph (3-5 sentences) with the headline conclusion and key investment implication.',
    'summary': 'Write 2-3 paragraphs covering key findings and investment implications. Be concise but substantive.',
    'standard': 'Write a 1-2 page report balancing positives and negatives with a clear investment conclusion. Include key data points.',
    'deep': 'Write a detailed 3-4 page analysis with data tables, bull/bear cases, and a comprehensive investment conclusion.',
    'comprehensive': 'Write a comprehensive 5+ page deep-dive with full analysis, data appendices, competitive context, and detailed investment conclusion.',
}


def ensure_catalyst_folders():
    """Auto-scaffold CATALYSTS/{TICKER}/ folders for all universe tickers."""
    CATALYSTS_DIR.mkdir(parents=True, exist_ok=True)
    all_tickers = []
    for tickers in PIPELINE_SECTOR_MAP.values():
        all_tickers.extend(tickers)
    for ticker in all_tickers:
        (CATALYSTS_DIR / ticker).mkdir(exist_ok=True)
    log.info(f"Catalyst folders scaffolded for {len(all_tickers)} tickers in {CATALYSTS_DIR}")
    # Also scaffold SUMMARIES folder for audio auto-processing
    SUMMARIES_DIR.mkdir(parents=True, exist_ok=True)
    (SUMMARIES_DIR / "Processed").mkdir(exist_ok=True)


def scan_catalyst_topics(ticker: str) -> list[dict]:
    """Scan CATALYSTS/{ticker}/ for topic subfolders and their file counts."""
    ticker_dir = CATALYSTS_DIR / ticker.upper()
    if not ticker_dir.exists():
        return []
    folders = []
    for item in sorted(ticker_dir.iterdir()):
        if item.is_dir() and not item.name.startswith('.'):
            files = [f.name for f in item.iterdir() if f.is_file() and not f.name.startswith('.')]
            folders.append({
                'name': item.name,
                'fileCount': len(files),
                'files': files[:50],  # Cap at 50 filenames
            })
    return folders


_known_files = {}  # key: folder path, value: set of filenames

def check_for_new_files():
    """Scan CATALYSTS and STOCKS folders for new files, create alerts for any found."""
    global _known_files
    SOURCE_EXTS = {'.pdf', '.xlsx', '.xls', '.csv', '.txt', '.md', '.docx'}
    alerts_created = 0

    for base_dir, folder_type in [(CATALYSTS_DIR, 'catalysts'), (STOCKS_DIR, 'stocks')]:
        if not base_dir.exists():
            continue
        try:
            for ticker_dir in sorted(base_dir.iterdir()):
                if not ticker_dir.is_dir() or ticker_dir.name.startswith('.'):
                    continue
                ticker = ticker_dir.name.upper()

                if folder_type == 'catalysts':
                    # Scan topic subfolders
                    for topic_dir in sorted(ticker_dir.iterdir()):
                        if not topic_dir.is_dir() or topic_dir.name.startswith('.'):
                            continue
                        folder_key = str(topic_dir)
                        current_files = set()
                        for f in topic_dir.iterdir():
                            if f.is_file() and not f.name.startswith('.') and f.suffix.lower() in SOURCE_EXTS:
                                current_files.add(f.name)

                        if folder_key not in _known_files:
                            _known_files[folder_key] = current_files
                            continue  # First scan — don't alert on existing files

                        new_files = current_files - _known_files[folder_key]
                        if new_files:
                            _known_files[folder_key] = current_files
                            try:
                                requests.post(
                                    f"{CHARLIE_API}/api/alerts",
                                    json={
                                        'alertType': 'new_files',
                                        'ticker': ticker,
                                        'title': f"{len(new_files)} new file{'s' if len(new_files) > 1 else ''} in {ticker}/{topic_dir.name}",
                                        'detail': {
                                            'folder': f'CATALYSTS/{ticker}/{topic_dir.name}',
                                            'folderType': 'catalysts',
                                            'topic': topic_dir.name,
                                            'newFiles': sorted(new_files),
                                            'totalFiles': len(current_files),
                                        },
                                    },
                                    headers=_agent_headers(),
                                    timeout=10,
                                )
                                alerts_created += 1
                            except Exception as e:
                                log.debug(f"Alert creation failed: {e}")
                else:
                    # Scan main STOCKS folder (not Processed/ or Prior Versions/)
                    folder_key = str(ticker_dir)
                    current_files = set()
                    for f in ticker_dir.iterdir():
                        if f.is_file() and not f.name.startswith('.') and f.suffix.lower() in SOURCE_EXTS:
                            current_files.add(f.name)

                    if folder_key not in _known_files:
                        _known_files[folder_key] = current_files
                        continue

                    new_files = current_files - _known_files[folder_key]
                    if new_files:
                        _known_files[folder_key] = current_files
                        try:
                            requests.post(
                                f"{CHARLIE_API}/api/alerts",
                                json={
                                    'alertType': 'new_files',
                                    'ticker': ticker,
                                    'title': f"{len(new_files)} new file{'s' if len(new_files) > 1 else ''} in {ticker}",
                                    'detail': {
                                        'folder': f'STOCKS/{ticker}',
                                        'folderType': 'stocks',
                                        'newFiles': sorted(new_files),
                                        'totalFiles': len(current_files),
                                    },
                                },
                                headers=_agent_headers(),
                                timeout=10,
                            )
                            alerts_created += 1
                        except Exception as e:
                            log.debug(f"Alert creation failed: {e}")
        except PermissionError:
            pass  # iCloud folder not accessible
    if alerts_created:
        log.info(f"Created {alerts_created} new file alert(s)")


_known_audio_files = set()
_audio_initialized = False
AUDIO_EXTS = {'.mp3', '.mp4', '.m4a', '.wav', '.webm', '.ogg', '.flac'}

def check_for_new_audio():
    """Scan SUMMARIES/ folder for new audio files, auto-process them."""
    global _known_audio_files, _audio_initialized
    if not SUMMARIES_DIR.exists():
        return
    try:
        current_files = set()
        for f in SUMMARIES_DIR.iterdir():
            if f.is_file() and not f.name.startswith('.') and f.suffix.lower() in AUDIO_EXTS:
                current_files.add(f.name)

        # On first tick, treat files already present as NEW so pre-existing
        # audio gets processed. Processed files move to SUMMARIES/Processed/
        # so restart doesn't re-trigger them.
        if not _audio_initialized:
            new_files = current_files
            _known_audio_files = current_files
            _audio_initialized = True
            if not new_files:
                return
        else:
            new_files = current_files - _known_audio_files
            if not new_files:
                return
            _known_audio_files = current_files
        for fname in new_files:
            fpath = SUMMARIES_DIR / fname
            log.info(f"New audio file detected: {fname}, auto-processing...")
            try:
                with open(fpath, 'rb') as af:
                    file_data = af.read()
                # Upload to backend auto-process endpoint
                res = requests.post(
                    f"{CHARLIE_API}/api/auto-process-audio",
                    files={'file': (fname, file_data)},
                    data={'detailLevel': 'standard'},
                    headers={'Authorization': f'ApiKey {os.environ.get("CHARLIE_API_KEY", "")}'},
                    timeout=30,
                )
                if res.ok:
                    data = res.json()
                    log.info(f"Audio auto-process started: {fname} (job {data.get('jobId', '?')})")
                    # Move to Processed/
                    processed_dir = SUMMARIES_DIR / "Processed"
                    processed_dir.mkdir(exist_ok=True)
                    try:
                        fpath.rename(processed_dir / fname)
                        log.info(f"Moved {fname} to Processed/")
                    except Exception as e:
                        log.warning(f"Could not move {fname} to Processed/: {e}")
                else:
                    log.warning(f"Audio auto-process failed for {fname}: {res.status_code}")
            except Exception as e:
                log.warning(f"Error auto-processing audio {fname}: {e}")
    except PermissionError:
        pass


# ------------------------------------------------------------------------------
# AUTOMATED CATALYST SYNTHESIS
# Watches CATALYSTS/{TICKER}/{topic}/ folders and auto-fires a synthesis job
# when docs are added/updated, with debounce + per-topic cooldown.
# ------------------------------------------------------------------------------

CATALYST_AUTO_STATE_FILE = Path.home() / "Library/Application Support/charlie-agent/auto-catalyst-state.json"
CATALYST_AUTO_DEBOUNCE_S = 120     # wait this long after last file change before firing
CATALYST_AUTO_COOLDOWN_S = 15 * 60 # don't re-fire same topic within this window
CATALYST_SOURCE_EXTS = {'.pdf', '.xlsx', '.xls', '.csv', '.txt', '.md', '.docx', '.doc', '.pptx', '.ppt', '.html'}


def _load_catalyst_auto_state() -> dict:
    try:
        if CATALYST_AUTO_STATE_FILE.exists():
            return json.loads(CATALYST_AUTO_STATE_FILE.read_text())
    except Exception as e:
        log.debug(f"auto-catalyst state load failed: {e}")
    return {}


def _save_catalyst_auto_state(state: dict) -> None:
    try:
        CATALYST_AUTO_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
        CATALYST_AUTO_STATE_FILE.write_text(json.dumps(state, indent=2))
    except Exception as e:
        log.debug(f"auto-catalyst state save failed: {e}")


_catalyst_auto_mode_cache = {'checked_at': 0.0, 'value': {'enabled': False, 'expires_at': None}}

def _fetch_catalyst_auto_mode() -> dict:
    """Cached fetch of /api/catalysts/auto-mode. Refreshes every 30s."""
    now_t = time.time()
    if now_t - _catalyst_auto_mode_cache['checked_at'] < 30:
        return _catalyst_auto_mode_cache['value']
    try:
        r = requests.get(f"{CHARLIE_API}/api/catalysts/auto-mode",
                         headers=_agent_headers(), timeout=8)
        if r.ok:
            _catalyst_auto_mode_cache['value'] = r.json()
            _catalyst_auto_mode_cache['checked_at'] = now_t
    except Exception as e:
        log.debug(f"auto-mode fetch failed: {e}")
    return _catalyst_auto_mode_cache['value']


def check_for_catalyst_auto_synth() -> None:
    """Scan CATALYSTS/{TICKER}/{topic}/ folders and auto-fire synthesis jobs
    when files change. Uses a debounce + cooldown to avoid firing while the
    user is still dropping files, and to avoid duplicate synths."""
    if not CATALYSTS_DIR.exists():
        return
    now = time.time()
    state = _load_catalyst_auto_state()

    try:
        for ticker_dir in CATALYSTS_DIR.iterdir():
            if not ticker_dir.is_dir() or ticker_dir.name.startswith('.'):
                continue
            ticker = ticker_dir.name.upper()

            for topic_dir in ticker_dir.iterdir():
                if not topic_dir.is_dir() or topic_dir.name.startswith('.') or topic_dir.name.lower() == 'processed':
                    continue
                topic = topic_dir.name

                # Collect source docs in this topic folder
                source_files = []
                max_mtime = 0.0
                try:
                    for f in topic_dir.iterdir():
                        if f.is_file() and not f.name.startswith('.') and f.suffix.lower() in CATALYST_SOURCE_EXTS:
                            source_files.append(f.name)
                            try:
                                max_mtime = max(max_mtime, f.stat().st_mtime)
                            except Exception:
                                pass
                except PermissionError:
                    continue

                if not source_files:
                    continue

                key = f"{ticker}::{topic}"
                entry = state.get(key, {})
                prior_mtime = entry.get('last_seen_mtime', 0.0)
                prior_fired = entry.get('last_fired_at', 0.0)
                prior_fingerprint = entry.get('fingerprint', '')
                fingerprint = f"{len(source_files)}@{int(max_mtime)}"

                # Keep state fresh so we don't lose track of what's there
                state[key] = {
                    'last_seen_mtime': max_mtime,
                    'last_fired_at': prior_fired,
                    'fingerprint': fingerprint,
                    'last_file_count': len(source_files),
                }

                # Skip if fingerprint hasn't changed AND we've already fired for it
                if fingerprint == prior_fingerprint and prior_fired > 0:
                    continue

                # Debounce: wait until files have been quiet for DEBOUNCE_S
                if now - max_mtime < CATALYST_AUTO_DEBOUNCE_S:
                    continue

                # Cooldown: don't re-fire same topic within COOLDOWN_S
                if now - prior_fired < CATALYST_AUTO_COOLDOWN_S:
                    continue

                # Decide: auto-fire or just propose?
                auto_mode = _fetch_catalyst_auto_mode()
                endpoint = 'synthesize' if auto_mode.get('enabled') else 'propose-synth'
                action_word = 'firing' if auto_mode.get('enabled') else 'proposing'
                log.info(f"Auto-catalyst: {action_word} {ticker}/{topic} ({len(source_files)} files)")
                try:
                    res = requests.post(
                        f"{CHARLIE_API}/api/catalysts/{endpoint}",
                        json={
                            'ticker': ticker,
                            'topic': topic,
                            'length': 'standard',
                            'customInstructions': '',
                            'fingerprint': fingerprint,
                            'fileCount': len(source_files),
                        },
                        headers=_agent_headers(),
                        timeout=15,
                    )
                    if res.ok:
                        job_id = res.json().get('jobId', '?')
                        log.info(f"Auto-catalyst: {action_word} done for {ticker}/{topic} (job {job_id})")
                        state[key]['last_fired_at'] = now
                        state[key]['fingerprint'] = fingerprint
                    else:
                        log.warning(f"Auto-catalyst POST failed for {ticker}/{topic}: {res.status_code} {res.text[:200]}")
                except Exception as e:
                    log.warning(f"Auto-catalyst error for {ticker}/{topic}: {e}")

        _save_catalyst_auto_state(state)
    except Exception as e:
        log.debug(f"Auto-catalyst scan error: {e}")


def process_scan_catalysts_job(job: dict) -> None:
    """Scan catalyst folders and report back to the backend."""
    job_id = job.get("id", "")
    ticker = job.get("ticker", "").upper()
    try:
        ensure_catalyst_folders()

        if ticker == "ALL":
            # Scan all tickers
            all_tickers = []
            for tickers in PIPELINE_SECTOR_MAP.values():
                all_tickers.extend(tickers)
            for tk in all_tickers:
                folders = scan_catalyst_topics(tk)
                requests.post(
                    f"{CHARLIE_API}/api/catalysts/folders",
                    json={"ticker": tk, "folders": folders},
                    headers=_agent_headers(),
                    timeout=15,
                )
        else:
            folders = scan_catalyst_topics(ticker)
            requests.post(
                f"{CHARLIE_API}/api/catalysts/folders",
                json={"ticker": ticker, "folders": folders},
                headers=_agent_headers(),
                timeout=15,
            )

        update_job_progress(job_id, "complete", "Scan complete", 100)
        log.info(f"Catalyst folder scan complete for {ticker}")
    except Exception as e:
        log.error(f"Catalyst scan failed: {e}")
        update_job_progress(job_id, "failed", "Failed", None, error=str(e))


def _generate_source_provenance_local(client, source_names, synthesis_markdown, ticker, topic):
    """Generate source provenance summary using Claude directly (local agent path)."""
    if not source_names:
        return None
    try:
        file_list = '\n'.join([f"- {name}" for name in source_names])
        prompt = f"""You just produced a synthesis report for {ticker} on the topic "{topic}". Below are the source documents you had access to and the synthesis you wrote.

SOURCE DOCUMENTS PROVIDED:
{file_list}

SYNTHESIS OUTPUT (first 3000 chars):
{synthesis_markdown[:3000]}

Now write a brief SOURCE PROVENANCE summary for the analyst's own reference. For each source document:
1. State the document name
2. Describe in 1-2 sentences what key information or data points from that document were used in the synthesis
3. If a document was not materially used (e.g., duplicate content, irrelevant), note that

Keep it concise and factual. Use markdown formatting with bullet points. This is an internal reference, not part of the analysis."""

        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
            system="You are an analyst documenting your research process. Be specific about which data points came from which sources.",
            messages=[{"role": "user", "content": prompt}],
        )
        return response.content[0].text
    except Exception as e:
        log.warning(f"Source provenance generation failed: {e}")
        return f"Source provenance generation failed. {len(source_names)} files were provided: {', '.join(source_names)}"


def process_synthesis_job(job: dict, api_key: str) -> None:
    """Process a catalyst synthesis job -- read local files, call Claude, upload result."""
    job_id = job.get("id", "")
    ticker = job.get("ticker", "").upper()
    steps_detail = job.get("steps_detail", {})
    if isinstance(steps_detail, str):
        try:
            steps_detail = json.loads(steps_detail)
        except Exception:
            steps_detail = {}

    topic = steps_detail.get("topic", "")
    length = steps_detail.get("length", "standard")
    custom_instructions = steps_detail.get("customInstructions", "")

    try:
        update_job_progress(job_id, "running", "Reading source files...", 10)
        notify(f"*Charlie Agent:* Starting synthesis for {ticker}/{topic}")

        # Read files from CATALYSTS/{ticker}/{topic}/
        topic_dir = CATALYSTS_DIR / ticker / topic
        if not topic_dir.exists():
            raise FileNotFoundError(f"Folder not found: {topic_dir}")

        excluded_files = set(steps_detail.get('excludedFiles', []))
        source_parts = []
        file_count = 0
        for fpath in sorted(topic_dir.iterdir()):
            if fpath.is_file() and not fpath.name.startswith('.') and fpath.name not in excluded_files:
                file_count += 1
                ext = fpath.suffix.lower()
                try:
                    if ext == '.pdf':
                        # Read PDF as base64 for Claude's native PDF support
                        pdf_b64 = base64.b64encode(fpath.read_bytes()).decode('ascii')
                        source_parts.append({
                            'type': 'pdf',
                            'name': fpath.name,
                            'data': pdf_b64,
                        })
                    elif ext in ('.docx',):
                        # Extract text from docx
                        try:
                            from docx import Document as DocxDocument
                            doc = DocxDocument(str(fpath))
                            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                            source_parts.append({
                                'type': 'text',
                                'name': fpath.name,
                                'content': text,
                            })
                        except Exception as e2:
                            log.warning(f"Could not read docx {fpath.name}: {e2}")
                            source_parts.append({'type': 'text', 'name': fpath.name, 'content': f'[Could not read: {e2}]'})
                    elif ext in ('.xlsx', '.xls'):
                        # Extract via zipfile+XML or openpyxl
                        try:
                            import openpyxl
                            wb = openpyxl.load_workbook(str(fpath), data_only=True)
                            text_parts = []
                            for ws in wb.worksheets:
                                text_parts.append(f"Sheet: {ws.title}")
                                for row in ws.iter_rows(values_only=True):
                                    vals = [str(c) if c is not None else '' for c in row]
                                    if any(vals):
                                        text_parts.append('\t'.join(vals))
                            source_parts.append({'type': 'text', 'name': fpath.name, 'content': '\n'.join(text_parts)})
                        except Exception as e2:
                            log.warning(f"Could not read xlsx {fpath.name}: {e2}")
                    elif ext in ('.txt', '.md', '.csv', '.tsv'):
                        text = fpath.read_text(errors='replace')
                        source_parts.append({'type': 'text', 'name': fpath.name, 'content': text})
                    else:
                        log.debug(f"Skipping unsupported file type: {fpath.name}")
                except Exception as e2:
                    log.warning(f"Error reading {fpath.name}: {e2}")

        if not source_parts:
            raise ValueError(f"No readable files found in {topic_dir}")

        log.info(f"Read {file_count} files from {topic_dir}")
        update_job_progress(job_id, "running", f"Read {file_count} files, calling Claude...", 30)

        # Build Claude API call
        client = anthropic.Anthropic(api_key=api_key)
        length_instruction = CATALYST_LENGTH_PRESETS.get(length, CATALYST_LENGTH_PRESETS['standard'])
        custom_block = f"\nAdditional Instructions: {custom_instructions}" if custom_instructions else ""

        # Estimate tokens and split into batches if needed
        def _estimate_source_tokens(parts):
            total = 0
            for sp in parts:
                if sp['type'] == 'pdf':
                    total += int(len(sp['data']) * 0.3)  # base64 -> tokens rough estimate
                else:
                    total += int(len(sp.get('content', '')) * 0.3)
            return total

        def _split_source_batches(parts, max_tokens=170_000):
            batches, current, current_tok = [], [], 0
            for sp in parts:
                est = _estimate_source_tokens([sp])
                if current_tok + est > max_tokens and current:
                    batches.append(current)
                    current, current_tok = [], 0
                current.append(sp)
                current_tok += est
            if current:
                batches.append(current)
            return batches if batches else [parts]

        est_tokens = _estimate_source_tokens(source_parts)
        batches = _split_source_batches(source_parts)
        total_batches = len(batches)

        if total_batches > 1:
            log.info(f"Sources exceed context limit (~{est_tokens:,} est tokens). Splitting into {total_batches} batches.")

        def _build_content_blocks(parts, prompt_text):
            blocks = []
            pdf_count = 0
            for sp in parts:
                if sp['type'] == 'pdf':
                    block = {
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": sp['data']},
                    }
                    if pdf_count < 4:
                        block["cache_control"] = {"type": "ephemeral"}
                    blocks.append(block)
                    blocks.append({"type": "text", "text": f"[Document: {sp['name']}]"})
                    pdf_count += 1
            text_sources = '\n\n---\n\n'.join([
                f"### Source: {sp['name']}\n{sp['content'][:8000]}"
                for sp in parts if sp['type'] == 'text'
            ])
            if text_sources:
                blocks.append({"type": "text", "text": text_sources})
            blocks.append({"type": "text", "text": prompt_text})
            return blocks

        # Process first batch
        update_job_progress(job_id, "running", f"Synthesizing report (batch 1/{total_batches})...", 50)

        prompt_text = CATALYST_SYNTHESIS_PROMPT.format(
            length_instruction=length_instruction,
            ticker=ticker,
            topic=topic,
            custom_instructions=custom_block,
            source_content="[See attached documents above]",
        )
        content_blocks = _build_content_blocks(batches[0], prompt_text)

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8192,
            system="You are a senior equity research analyst. Follow all instructions precisely.",
            messages=[{"role": "user", "content": content_blocks}],
        )
        markdown = response.content[0].text

        # Process subsequent batches — merge into existing synthesis
        for i, batch in enumerate(batches[1:], 2):
            update_job_progress(job_id, "running", f"Synthesizing report (batch {i}/{total_batches})...", 50 + int(30 * i / total_batches))
            merge_prompt = f"""You previously wrote a synthesis report based on earlier documents. Now incorporate these ADDITIONAL source documents into the existing report.

EXISTING REPORT:
{markdown[:6000]}

## ATTRIBUTION RULES — ABSOLUTE, NO EXCEPTIONS
- ZERO references to any broker, bank, sellside firm, or analyst by name. Not "Wolfe Research", not "Goldman Sachs", not "JPMorgan", not ANY firm.
- Replace ALL firm attributions with first-person voice: "I estimate", "my analysis shows".
- Write ENTIRELY in FIRST PERSON as YOUR OWN proprietary work.
- If the existing report accidentally contains firm names, REMOVE them in this pass.

## MERGE INSTRUCTIONS
- Integrate new findings into the existing report structure
- Give EQUAL WEIGHT to new documents — do not let earlier batch dominate
- Add new data points, update conclusions if warranted
- Keep the same format, tone, and length
- {length_instruction}
{custom_block}

Write the complete, updated synthesis report now. ZERO firm names, ALL first person."""
            merge_blocks = _build_content_blocks(batch, merge_prompt)
            merge_response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8192,
                system="You are a senior equity research analyst. Follow all instructions precisely.",
                messages=[{"role": "user", "content": merge_blocks}],
            )
            markdown = merge_response.content[0].text
            log.info(f"Batch {i}/{total_batches} merged: {len(markdown)} chars")

        log.info(f"Synthesis generated: {len(markdown)} chars")

        # Generate source provenance
        update_job_progress(job_id, "running", "Analyzing source contributions...", 80)
        source_names = [sp['name'] for sp in source_parts]
        provenance = _generate_source_provenance_local(client, source_names, markdown, ticker, topic)

        update_job_progress(job_id, "running", "Uploading results...", 90)

        # Upload result to backend
        result_data = {
            'markdown': markdown,
            'topic': topic,
            'length': length,
            'fileCount': file_count,
            'sourceFiles': source_names,
            'sourceProvenance': provenance,
        }

        requests.post(
            f"{CHARLIE_API}/api/pipeline/jobs/{job_id}/result",
            json={"status": "complete", "result": result_data},
            headers=_agent_headers(),
            timeout=30,
        )

        # Fallback: direct DB update via progress endpoint
        update_job_progress(job_id, "complete", "Complete", 100, result=result_data)
        notify(f"*Charlie Agent:* {ticker}/{topic} synthesis complete\n{file_count} docs, {len(markdown):,} chars")

    except Exception as e:
        log.error(f"Synthesis job {job_id[:12]}... failed: {e}")
        import traceback
        traceback.print_exc()
        update_job_progress(job_id, "failed", "Failed", None, error=str(e))
        notify(f"*Charlie Agent:* {ticker}/{topic} synthesis FAILED\n{str(e)[:200]}")


def process_ticker_directly(ticker: str, api_key: str, mode: str = "new") -> None:
    """Process a ticker directly without a backend job (--ticker flag)."""
    fake_job = {
        "id": str(uuid.uuid4()),
        "ticker": ticker.upper(),
        "job_type": "note",
        "steps_detail": json.dumps({"mode": mode}),
    }
    log.info(f"Direct processing mode for {ticker.upper()}")
    process_note_job(fake_job, api_key)


# ---------------------------------------------------------------------------
# Main loop
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Charlie Local Agent -- polls for note jobs and processes them locally"
    )
    parser.add_argument(
        "--api-key",
        help="Anthropic API key (or set ANTHROPIC_API_KEY env var)",
    )
    parser.add_argument(
        "--once",
        action="store_true",
        help="Process one job and exit",
    )
    parser.add_argument(
        "--ticker",
        help="Process a specific ticker directly (no backend job needed)",
    )
    parser.add_argument(
        "--mode",
        choices=["new", "update"],
        default="new",
        help="Note generation mode (default: new)",
    )
    parser.add_argument(
        "--debug",
        action="store_true",
        help="Enable debug logging",
    )
    args = parser.parse_args()

    # Resolve API key
    api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY") or load_api_key_from_config()
    if not api_key:
        print("Error: Anthropic API key required.")
        print("  Set ANTHROPIC_API_KEY env var, use --api-key, or save to ~/.charlie_agent_config.json")
        sys.exit(1)

    # Configure logging
    _log_path = os.path.expanduser("~/Library/Logs/charlie-agent.log")
    logging.basicConfig(
        level=logging.DEBUG if args.debug else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
        handlers=[logging.StreamHandler(), logging.FileHandler(_log_path)],
    )

    # Validate iCloud directory
    if not STOCKS_DIR.exists():
        log.error(f"iCloud STOCKS directory not found: {STOCKS_DIR}")
        sys.exit(1)

    # Scaffold catalyst folders on startup
    ensure_catalyst_folders()

    log.info("Charlie Local Agent starting")
    log.info(f"  Stocks dir: {STOCKS_DIR}")
    log.info(f"  Backend: {CHARLIE_API}")

    # Direct ticker mode
    if args.ticker:
        process_ticker_directly(args.ticker, api_key, args.mode)
        return

    # Polling mode
    shutdown = threading.Event()

    def handle_signal(signum, frame):
        log.info(f"Received signal {signum}, shutting down...")
        shutdown.set()

    signal.signal(signal.SIGINT, handle_signal)
    signal.signal(signal.SIGTERM, handle_signal)

    log.info(f"Polling for jobs every {POLL_INTERVAL}s (Ctrl+C to stop)")

    # One-time import of existing notes + documents from iCloud to Charlie
    try:
        import_existing_notes()
    except Exception as e:
        log.warning(f"Note import failed: {e}")
    try:
        import_existing_documents()
    except Exception as e:
        log.warning(f"Document import failed: {e}")

    consecutive_errors = 0
    MAX_CONSECUTIVE_ERRORS = 10
    MANIFEST_INTERVAL = 60  # seconds
    last_manifest_push = 0.0  # epoch; push immediately on first iteration

    while not shutdown.is_set():
        try:
            # Auto-unzip + push file manifest periodically (every 60s)
            now = time.time()
            if now - last_manifest_push >= MANIFEST_INTERVAL:
                try:
                    auto_unzip_stock_folders()
                    auto_unzip_catalyst_folders()
                    push_file_manifest()
                    process_pending_syncs()
                    check_for_new_files()
                    check_for_new_audio()
                    check_for_catalyst_auto_synth()
                    last_manifest_push = now
                    log.debug("File manifest pushed")
                except Exception as e:
                    log.debug(f"Manifest push error: {e}")
                    last_manifest_push = now  # don't retry immediately on error

            # Check for urgent document upload requests from pipeline (every cycle)
            try:
                check_and_fulfill_doc_requests()
            except Exception as e:
                log.debug(f"Doc request check error: {e}")

            jobs = poll_for_jobs()
            consecutive_errors = 0  # reset on successful poll

            if jobs:
                job = jobs[0]
                ticker = job.get("ticker", "???")
                log.info(f"Found job: {job['id'][:12]}... for {ticker}")

                # Claim the job
                claimed = claim_job(job["id"])
                if not claimed:
                    log.warning(f"Could not claim job {job['id'][:12]}..., skipping")
                    shutdown.wait(timeout=POLL_INTERVAL)
                    continue

                job_type = job.get("job_type", "note")
                if job_type == "synthesis":
                    process_synthesis_job(job, api_key)
                elif job_type == "scan_catalysts":
                    process_scan_catalysts_job(job)
                else:
                    process_note_job(job, api_key)

                if args.once:
                    break
            else:
                log.debug("No pending jobs")

        except KeyboardInterrupt:
            break
        except Exception as e:
            consecutive_errors += 1
            log.error(f"Unexpected error: {e}")
            if consecutive_errors >= MAX_CONSECUTIVE_ERRORS:
                log.error(f"Too many consecutive errors ({MAX_CONSECUTIVE_ERRORS}), exiting")
                sys.exit(1)

        shutdown.wait(timeout=POLL_INTERVAL)

    log.info("Agent stopped.")


if __name__ == "__main__":
    main()
