"""Document planning for research-note generation.

Extracted so the server can do what only the local agent could: decide which
documents go to the model, in what order, and in which batch.

The ordering is the point. call_claude() on the agent generates the note from
the FIRST batch and merely merges later batches into it, so whichever documents
land in batch one decide the note's spine. Those batches were filled in
whatever order the filesystem returned, which meant a broker PDF could shape
the thesis while the 10-K arrived as an afterthought. Ranking first makes that
deliberate.
"""

import base64
import io
import math
import re
from datetime import datetime, timezone

# Tokens per byte of PDF.
#
# The agent uses 0.4, which cannot be right: it implies a 1MB PDF costs 400,000
# tokens, so a 170K budget would fit only ~425KB and nearly every document would
# land in a batch of its own. Since batch one writes the note and later batches
# are only merged into it, that quietly reduced most multi-document notes to
# "one document, plus footnotes".
#
# A text PDF runs roughly 30-50 pages per MB at ~1-3K tokens a page, so ~0.1
# is closer, and image-heavy decks are lower still. Estimating a little low is
# the safer error: an over-full batch returns a context-length error we catch
# and re-split, whereas over-splitting silently degrades the note and nothing
# reports it.
APPROX_TOKENS_PER_BYTE = 0.12
MAX_TOKENS_PER_BATCH = 170_000        # leaves room for prompt + response under 200K

# A PDF larger than this is sent as extracted text rather than as a document
# block. Base64 inflates a file by about a third and the model has a hard
# per-request ceiling, so a large PDF attached whole fails at the API rather
# than at upload -- raising a size cap alone would just move the error.
PDF_INLINE_LIMIT_BYTES = 12 * 1024 * 1024

# What a document is worth to the note's spine. Primary sources describe the
# business; broker material interprets it. When only some documents fit in the
# first batch, these should be the ones that do.
_KIND_RANK = [
    (r"10[-_ ]?k|annual[ _-]?report", 100, "annual report"),
    (r"10[-_ ]?q|quarterly[ _-]?report", 90, "quarterly report"),
    (r"transcript|earnings[ _-]?call|prepared[ _-]?remarks", 85, "transcript"),
    (r"investor[ _-]?day|analyst[ _-]?day|capital[ _-]?markets[ _-]?day", 80, "investor day"),
    (r"presentation|deck|slides|supplement", 70, "company presentation"),
    (r"press[ _-]?release|8[-_ ]?k|guidance", 65, "press release"),
    (r"proxy|def[ _-]?14a", 50, "proxy"),
    (r"model|forecast|\.xlsx?$|\.csv$", 45, "model"),
    (r"initiation|note|research|report", 30, "broker research"),
]
_DEFAULT_RANK = 40


def classify_document(filename):
    """(priority, human label) for a filename."""
    name = (filename or "").lower()
    for pattern, score, label in _KIND_RANK:
        if re.search(pattern, name):
            return score, label
    return _DEFAULT_RANK, "document"


def _recency_bonus(created_at):
    """Newer documents describe the company as it is now. Worth up to ~25."""
    if not created_at:
        return 0
    try:
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
        if created_at.tzinfo is None:
            created_at = created_at.replace(tzinfo=timezone.utc)
        age_days = (datetime.now(timezone.utc) - created_at).days
    except Exception:
        return 0
    if age_days <= 0:
        return 25
    # Half the bonus every ~180 days.
    return int(25 * math.exp(-age_days / 260.0))


def rank_documents(docs):
    """Order documents by how much they should shape the note.

    Returns the same dicts with `priority`, `kind` and `est_tokens` added, most
    important first. Ties keep their original order so the result is stable.
    """
    ranked = []
    for i, d in enumerate(docs or []):
        score, label = classify_document(d.get("filename"))
        priority = score + _recency_bonus(d.get("created_at") or d.get("createdAt"))
        out = dict(d)
        out["priority"] = priority
        out["kind"] = label
        # prepare_documents may already have costed this document (extracted
        # text is far cheaper than the pages it came from); do not undo that.
        out["est_tokens"] = d.get("est_tokens") or estimate_tokens([d])
        ranked.append((priority, -i, out))
    ranked.sort(key=lambda t: (-t[0], -t[1]))
    return [r[2] for r in ranked]


TOKENS_PER_PDF_PAGE = 2_000     # text plus the page image Anthropic renders


def pdf_page_count(doc):
    """Pages in a stored PDF, or None if it cannot be read.

    Pages are what a PDF actually costs. Bytes are not: a scanned page can be
    500KB and a text page 20KB while both cost about the same to read, so any
    size-based estimate is wrong by an order of magnitude in one direction or
    the other depending on the document. Cached on the dict so a planning pass
    and a run pass do not both parse the file.
    """
    if doc.get("_pages") is not None:
        return doc["_pages"] or None
    raw = doc.get("file_data")
    if not raw:
        return None
    try:
        from PyPDF2 import PdfReader
        data = base64.b64decode(raw) if isinstance(raw, str) else raw
        pages = len(PdfReader(io.BytesIO(data)).pages)
    except Exception:
        pages = 0
    doc["_pages"] = pages
    return pages or None


def estimate_tokens(docs):
    """Token cost of a document set, by pages where the file can be read.

    An estimate already attached to a document wins: prepare_documents may have
    decided to send a large PDF as extracted text, which costs a fraction of
    its pages, and re-deriving from the page count would throw that away.
    """
    total = 0
    for d in docs or []:
        if d.get("est_tokens"):
            total += d["est_tokens"]
            continue
        size = d.get("file_size") or d.get("fileSize") or 0
        if not size and d.get("file_data"):
            size = int(len(d["file_data"]) * 0.75)      # base64 -> bytes
        is_pdf = ("pdf" in (d.get("file_type") or d.get("fileType") or "").lower()
                  or str(d.get("filename", "")).lower().endswith(".pdf"))
        if is_pdf:
            pages = pdf_page_count(d)
            total += (pages * TOKENS_PER_PDF_PAGE if pages
                      else int(size * APPROX_TOKENS_PER_BYTE))
        else:
            name = str(d.get("filename", "")).lower()
            if name.endswith((".xlsx", ".xls", ".docx", ".pptx")):
                # Zipped XML: the readable text is a small fraction of the file.
                total += int(size * 0.04)
            else:
                # Plain text and CSV: roughly four bytes per token.
                total += int(size * 0.25)
    return total


def extract_pdf_text(doc, max_tokens=MAX_TOKENS_PER_BATCH // 2):
    """Text of a PDF, for documents too large to attach as pages.

    A 160-page 10-K costs more than a whole batch as page images, and it is the
    single most important document for the note -- so it must not be the one we
    drop. As text it costs roughly half as much and loses only layout, which a
    10-K barely uses. If even the text is too long, the front and back are kept:
    the business description and MD&A open the filing and the financials close
    it, while the middle is largely boilerplate.
    """
    raw = doc.get("file_data")
    if not raw:
        return ""
    try:
        from PyPDF2 import PdfReader
        data = base64.b64decode(raw) if isinstance(raw, str) else raw
        reader = PdfReader(io.BytesIO(data))
        pages = [(p.extract_text() or "") for p in reader.pages]
    except Exception:
        return ""

    budget_chars = max_tokens * 4
    joined = "\n\n".join(pages)
    if len(joined) <= budget_chars:
        return joined

    head = int(budget_chars * 0.65)
    tail = budget_chars - head
    return (joined[:head]
            + "\n\n[... middle of document omitted to fit context ...]\n\n"
            + joined[-tail:])


def extract_file_text(doc, max_chars=120_000):
    """Readable text from a spreadsheet, Word file, or plain text document.

    These were marked send_as="file" and then silently dropped: the content
    builder only ever emitted blocks for PDFs, so a broker model spreadsheet was
    ranked, charged against the token budget, reported as included in the plan,
    and never actually sent. Forward-year segment estimates usually live in
    exactly those files, so the note could only ever see what the PDFs happened
    to state.

    Spreadsheets are rendered per sheet as tab-separated rows with their values
    (not formulas), which is what a model can read. Empty rows and columns are
    dropped so a sparse sheet does not spend the budget on blanks. Returns ''
    when nothing can be read -- the caller then skips the document rather than
    sending an empty block.
    """
    import base64

    name = str(doc.get("filename", "")).lower()
    raw = doc.get("file_data") or doc.get("fileData") or ""
    if not raw:
        return ""
    try:
        blob = base64.b64decode(raw)
    except Exception:
        return ""

    parts = []
    try:
        if name.endswith((".xlsx", ".xlsm", ".xls")):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(blob), data_only=True, read_only=True)
            for ws in wb.worksheets:
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v).strip() for v in row]
                    while cells and not cells[-1]:
                        cells.pop()
                    if any(cells):
                        rows.append("\t".join(cells))
                if rows:
                    parts.append(f"--- sheet: {ws.title} ---")
                    parts.extend(rows)
                if sum(len(x) for x in parts) > max_chars:
                    break
            try:
                wb.close()
            except Exception:
                pass
        elif name.endswith((".docx", ".doc")):
            import docx
            d = docx.Document(io.BytesIO(blob))
            parts.extend(p.text for p in d.paragraphs if p.text.strip())
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append("\t".join(cells))
        elif name.endswith((".csv", ".txt", ".md", ".json")):
            parts.append(blob.decode("utf-8", errors="replace"))
    except Exception as e:
        print(f"[notegen] could not read {doc.get('filename')}: "
              f"{type(e).__name__}: {e}")
        return ""

    text = "\n".join(parts).strip()
    return text[:max_chars]


def prepare_documents(docs, max_tokens=MAX_TOKENS_PER_BATCH):
    """Decide how each document is sent, and mark it.

    `send_as` is "pages" for a normal PDF, "text" for one too large to attach,
    "file" for spreadsheets and plain text. Set before batching so the token
    estimate reflects what will actually be sent.
    """
    out = []
    for d in docs or []:
        d = dict(d)
        name = str(d.get("filename", "")).lower()
        is_pdf = ("pdf" in (d.get("file_type") or d.get("fileType") or "").lower()
                  or name.endswith(".pdf"))
        if not is_pdf:
            # Extract now, so the estimate below counts the text that will
            # actually be sent rather than the zipped file size.
            text = extract_file_text(d)
            if text:
                d["send_as"] = "file"
                d["extracted_text"] = text
                d["est_tokens"] = len(text) // 4
            else:
                d["send_as"] = "skip"
                d["skip_reason"] = "no readable text could be extracted"
        elif estimate_tokens([d]) > max_tokens:
            text = extract_pdf_text(d, max_tokens=max_tokens // 2)
            if text:
                d["send_as"] = "text"
                d["extracted_text"] = text
                d["est_tokens"] = len(text) // 4
            else:
                # Unreadable and oversized: attaching it would fail the request.
                d["send_as"] = "skip"
                d["skip_reason"] = "too large to attach and text could not be extracted"
        else:
            d["send_as"] = "pages"
        out.append(d)
    return out


def plan_batches(docs, max_tokens=MAX_TOKENS_PER_BATCH):
    """Rank, then split into batches that fit the context window.

    Batch one carries the highest-priority documents because it is the batch
    that writes the note.
    """
    ranked = rank_documents(prepare_documents(docs, max_tokens))
    ranked = [d for d in ranked if d.get("send_as") != "skip"]
    batches, current, current_tokens = [], [], 0
    for d in ranked:
        est = d.get("est_tokens") or estimate_tokens([d])
        if current and current_tokens + est > max_tokens:
            batches.append(current)
            current, current_tokens = [], 0
        current.append(d)
        current_tokens += est
    if current:
        batches.append(current)
    return batches or ([ranked] if ranked else [])


def plan_summary(docs, max_tokens=MAX_TOKENS_PER_BATCH):
    """What the run will do, for the user to see before committing to it.

    The agent logged this and threw it away; the person choosing the documents
    is the one who needs it.
    """
    batches = plan_batches(docs, max_tokens)
    total_tokens = sum(estimate_tokens(b) for b in batches)
    prepared = prepare_documents(docs, max_tokens)
    oversize = [d.get("filename") for d in prepared if d.get("send_as") == "text"]
    skipped = [{"filename": d.get("filename"), "reason": d.get("skip_reason")}
               for d in prepared if d.get("send_as") == "skip"]
    return {
        "documentCount": len(docs or []),
        "batchCount": len(batches),
        "estimatedTokens": total_tokens,
        # Input pricing dominates; output is a few thousand tokens per batch.
        "estimatedCostUsd": round(total_tokens / 1_000_000 * 3.0, 2),
        "estimatedMinutes": max(1, round(len(batches) * 1.5 + total_tokens / 120_000)),
        "sentAsText": oversize,
        "skipped": skipped,
        "batches": [
            {
                "index": i + 1,
                "shapesTheNote": i == 0,
                "estimatedTokens": estimate_tokens(b),
                "documents": [
                    {"filename": d.get("filename"), "kind": d.get("kind"),
                     "priority": d.get("priority"), "sentAs": d.get("send_as"),
                     "estimatedTokens": d.get("est_tokens")}
                    for d in b
                ],
            }
            for i, b in enumerate(batches)
        ],
    }
